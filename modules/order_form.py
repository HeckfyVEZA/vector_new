from docx import Document as d
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
import modules.assitent_info #Просто обращаемся к тексту
import streamlit as st
import pandas as pd
from datetime import date
from modules.canal_blank_form import sizes_and_weight_recognition, megatable
#from import_google_table import table_vector

def rezerving(type_scheme):
    #st.write(type_scheme)
    match type_scheme:
        case '1':
            return [['Клапан обратный CVS16.05.020.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.020.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '2':
            return [['Клапан обратный CVS16.05.025.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.025.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '3':
            return [['Клапан обратный CVS16.05.025.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.025.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '4':
            return [['Клапан обратный CVS16.05.032.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.032.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '5':
            return [['Клапан обратный CVS16.05.032.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.032.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '6':
            return [['Клапан обратный CVS16.05.040.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.040.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '7':
            return [['Клапан обратный CVS16.05.050.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.050.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '8':
            return [['Клапан обратный CVS16.05.065.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.065.016.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '9':
            return [['Клапан обратный CVS16.05.080.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.080.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '10':
            return [['Клапан обратный CVS16.05.100.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.100.016.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '11':
            return [['Клапан обратный CVS16.05.125.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.125.016.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '1А':
            return [['Клапан обратный CVS16.05.020.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.020.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]
        case '1Б':
            return [['Клапан обратный CVS16.05.020.16 ТУ 3700-005-81673229-2009', 2], ['Кран шаровой фланцевый КШ.Ц.Ф.020.040.Н/П.02 ТУ 3742-001-45630744-2003', 4]]

def gabar_table(doc, vector, rezerv): # Заполняет таблицу с габаритами в шаблоне
    #st.write(vector)
    table = doc.tables[3]
    inf = megatable()["-".join(vector.split("-")[1:4])]
    table.rows[0].cells[-1].paragraphs[0].text = f"{str(inf[0])} (резьбовое)" if vector.split("-")[2] == "Ш" else f"ДУ {str(inf[0])} (фланцевое)"
    table.rows[1].cells[-1].paragraphs[0].text = {"П": "правая", "Л": "левая"}[vector.split("-")[4]]
    vecspli = vector.split("-")
    L = sizes_and_weight_recognition(str(vecspli[1]), str(vecspli[2]), str(vecspli[3]))["L"]
    Wei = sizes_and_weight_recognition(str(vecspli[1]), str(vecspli[2]), str(vecspli[3]))["weight"]
    h1 = sizes_and_weight_recognition(str(vecspli[1]), str(vecspli[2]), str(vecspli[3]))["H1"]
    table.rows[2].cells[-1].paragraphs[0].text = str(L) if not rezerv else " "
    table.rows[-1].cells[-1].paragraphs[0].text = str(Wei) if not rezerv else " "
    table.rows[-2].cells[-1].paragraphs[0].text = str(h1) if not rezerv else " "
    
    doc.tables[3].autofit
    return doc

def temps(valve:str, keyword): # Определяет температурные диапазоны для ВЕКТОРа
    match valve:
            case "С":
                temperatures = {"3": "0...+50 °C", "1": "+5...+150 °C","2": "+5...+150 °C", "4": "+5...+120 °C", "4М": "+5...+150 °C", "5": "+5...+120 °C", "5М": "+5...+120 °C", "6": "+5...+120 °C", "6М": "+5...+120 °C"}
            case "Ш":
                temperatures = {"2": "+5...+120 °C", "4": "+5...+120 °C", "4М": "+5...+150 °C", "5": "+5...+120 °C", "5М": "+5...+120 °C", "6": "+5...+120 °C", "6М": "+5...+120 °C"}
    return temperatures[keyword]


# doc = d('C:\\Users\\kushhov\\Desktop\\vector-main\\template.docx')
def fulfil_temp(cblank,type_scheme,Data_frame,rezerve,developer_name): # Сама функция вывода всего и вся в test.doc
    #st.write(type_scheme)
    #st.write(cblank)
    current_date = ".".join(str(date.today()).split("-")[::-1])
    doc = d('template_1.docx')
    #doc = d('./template.docx')
    # Таблица шапки
    doc.tables[0].rows[0].cells[0].paragraphs[0].text = f"Узел Регулирующий для бланк-заказа\n№{cblank['order form']} от {current_date}"
    doc.tables[1].rows[0].cells[1].paragraphs[0].text = cblank["object"]
    doc.tables[1].rows[1].cells[1].paragraphs[0].text = cblank["orderer"]
    doc.tables[1].rows[-1].cells[1].paragraphs[0].text = cblank["manager"]
    doc.tables[1].rows[-1].cells[-1].paragraphs[0].text =cblank["vector"]
    try:
        doc.tables[1].rows[-2].cells[1].paragraphs[0].text = cblank["system"]
    except:
        doc.tables[1].rows[-2].cells[1].paragraphs[0].text = '-'
    doc.tables[1].rows[-3].cells[1].paragraphs[0].text = cblank["order form"]
    doc.tables[1].rows[2].cells[-1].paragraphs[0].text = developer_name #Надо добавить
    doc.tables[1].autofit
    
    #st.write(doc.tables[0].style)
    doc.tables[2].rows[0].cells[1].text = f"На основании {int(cblank['glycol'])}% {cblank['glycol type'][:-1]}я" if cblank['glycol'] else modules.assitent_info.pure_water
    doc.tables[2].rows[0].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    #doc.tables[2].rows[1].cells[1].text = str(cblank['temperature']).replace(".",",")
      




    doc.tables[2].rows[1].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.tables[2].rows[2].cells[1].text = str(cblank['consumption_from_blank']).replace(".",",")
    doc.tables[2].rows[2].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    

    doc = gabar_table(doc, cblank["vector"], rezerve) 

    #doc.tables[3].rows[0].cells[0].paragraphs[0].add_run().add_picture('C:\\Users\\kushhov\\Desktop\\vector-main\\scheme.jpg', width=Mm(90)).alignment =  WD_ALIGN_PARAGRAPH.CENTER #, width=Inches(5), height=Inches(3)
    #doc.paragraphs[5].add_run().add_picture('C:\\Users\\kushhov\\Desktop\\vector-main\\\Scheme\\scheme.jpg', width=Mm(90)).alignment =  WD_ALIGN_PARAGRAPH.CENTER #, width=Inches(5), height=Inches(3)


    path ='2-С.bmp'

    if rezerve:
        doc_scheme = f"{type_scheme[0]}-{type_scheme[1].upper()}{'Р'}"
    else:
        doc_scheme = f"{type_scheme[0]}-{type_scheme[1].upper()}"
  
    
    #st.write(doc_scheme)

    
    match doc_scheme:
        case "1-С":
            path = '1-С.bmp'
        case "2-С":
            path = '2-С.bmp'
        case "2-СР":
            path = '2-СР.bmp'
        case "2-Ш":
            path = '2-Ш.bmp'
        case "2-ШР":
            path = '2-Ш.bmp'
        case "3-С":
            path = '3-С.bmp'
        case "3-СР":
            path = '3-С.bmp'
        case "4-С":
            path = '4-С.bmp'
        case "4-СР":
            path = '4-СР.bmp'
        case "4-Ш":
            path = '4-Ш.bmp'
        case "4-ШР":
            path = '4-Ш.bmp'
        case "4М-С":
            path = '4-С.bmp'
        case "4М-СР":
            path = '4-СР.bmp'
        case "5М-С":
            path = '5М-С.bmp'
        case "5М-СР":
            path = '5М-СР.bmp'
        case "5М-Ш":
            path = '5М-Ш.bmp'
        case "5М-ШР":
            path = '5М-Ш.bmp'
        case "5-С":
            path = '5-С.bmp'
        case "5-СР":
            path = '5-СР.bmp'
        case "5-Ш":
            path = '5-Ш.bmp'
        case "5-ШР":
            path = '5-Ш.bmp'
        case "6М-Ш":
            path = '6М-Ш.bmp'
        case "6М-ШР":
            path = '6М-Ш.bmp'
        case "6-Ш":
            path = '6-Ш.bmp'
        case "6-ШР":
            path = '6-Ш.bmp'

    #st.write(path)
    path ='Scheme/' + path
    #path = f"./{path}"
   

    #doc.paragraphs[58].add_run().add_picture(path, width=Mm(180)).alignment =  WD_ALIGN_PARAGRAPH.CENTER


    doc.tables[2].autofit
    #Data_frame = pd.read_excel('C:\\Users\\test\\Desktop\\vector_veza-master\\Table_vector.xlsx')
    #sheet_id = "1Qn-rGHE-mBaXHzVL9GHgGURV2OK0ZMRxcWerJ98qQhY"
    #Data_frame = pd.read_csv(f'https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv') неудачная попытка напрямую обрратиться к google таблице по ссылки
    
    #Data_frame = table_vector()

    #type_scheme = f"{type_scheme[0]}-{type_scheme[1]}-{type_scheme[2]}"
    #type_scheme = type_scheme.replace("С","C")
    Number_Vector = f"{type_scheme[0]}-{type_scheme[1]}-{type_scheme[2]}"
    Number_Vector = Number_Vector.upper().replace('С','C')
    K_Number_Vector = f"К-{Number_Vector}"

    n = 0
    for i in range(1,13):
        values_equipment = Data_frame[Number_Vector].values[i-1]
        count_equipment =  Data_frame[K_Number_Vector].values[i-1]
        if rezerve and (type_scheme[1] != 'Ш') and ('Насос' in values_equipment):
            count_equipment = 2
        doc.tables[4].rows[i].cells[1].text = str(values_equipment)
        doc.tables[4].rows[i].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[4].rows[i].cells[2].text = str(count_equipment)
        doc.tables[4].rows[i].cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        if values_equipment != '-':
            n +=1
    doc.tables[4].autofit
    
    if rezerve and type_scheme[1] != 'Ш':
        rez_list = rezerving(str(type_scheme[2]))
        #st.write(rez_list)
        values_equipment1 = rez_list[0][0]
        count_equipment1  = rez_list[0][1]
        values_equipment2 = rez_list[1][0]
        count_equipment2  = rez_list[1][1]
        doc.tables[4].rows[n+1].cells[1].text = str(values_equipment1)
        doc.tables[4].rows[n+1].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[4].rows[n+1].cells[2].text = str(count_equipment1)
        doc.tables[4].rows[n+1].cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[4].rows[n+2].cells[1].text = str(values_equipment2)
        doc.tables[4].rows[n+2].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.tables[4].rows[n+2].cells[2].text = str(count_equipment2)
        doc.tables[4].rows[n+2].cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        n +=2

    def remove_row(table, row): # Просто удаляет пустые строки
        tbl, tr = table._tbl, row._tr
        tbl.remove(tr)

    for _ in range(20-n):
        remove_row(doc.tables[4], doc.tables[4].rows[-1])


    #doc.add_picture('C:\\Users\\kushhov\\Desktop\\vector-main\\scheme.jpg')
    doc.paragraphs[-2].add_run().add_picture(path, width=Mm(90)).alignment =  WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-2].add_run().add_picture(f"Scheme/legend.bmp", width=Mm(60)) # БЛЯ 
    return doc

   # file_zip.write('C:\\Users\\kushhov\\Desktop\\vector-main\\test'+str(x+1)+'.docx', compress_type=zipfile.ZIP_DEFLATED)
   

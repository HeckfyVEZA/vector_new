"""Модуль, в котором я собрал разные функции, классы и прочие приблуды, которыми я пользуюсь в своих программах. Теперь можно его расширять, дополнять, улучшать, модернизировать вообще кому угодно - пишите сюда вообще всё, что хотите и считаете нужным

Итак, что тут есть?

Функции:
    checking_arguments - Проверка типа и значения переменной. Ещё может вывести длину переменной, если подобное возможо для данного типа переменой.

    file_names_csv_and_xlsx - Создание двух переменных-названий файлов с расширениями csv и xlsx.

    from_xlsx_to_csv - Конвертация xlsx в csv.

    from_csv_to_xlsx - Конвертация csv в xlsx .

    line_by_line - Запись строки в csv-файл в его конец.

    decision - Функция, которая возвращает значение в зависимости от вопроса.

    the_new_order - Позаимствованный из интернета код, который двигает туда-сюда листы в экселе.

    csv_cleaning - Функция очистки csv-файла - в целом, морально устарела.

    ideal_message - Печатает строку состояния работы цикла в виде "Сделано столько-то чего-то там, осталось столько-то чего-то там. Процентов: проценты. Прошло времени - столько-то секунд. Осталось столько-то секунд".

    dadata_inn_and_address - Функция, обращающаяся к Дадате и забирающая оттуда всю необходимую информацию по данному ИНН.

    dadata_left - Возвращает количество оставшихся запросов.

    current_database - Возвращает параметры базы данных в PostgreSQL.

    veza_design - Простенький метод, который задаёт созданное Майей оформление для программок.

    temporary_filename - Возвращает имя временного файла, который будет храниться в той же директории, что и программа.

    convert_using_win32 - Конвертация из различных форматов посредством модуля win32 (где конвертация - сохранение файла в нужном формате через заданное приложение)

    ordered_content_from_docx - Функция читает подряд docx-файл и возвращает все основные объекты файла в порядке их появления.

    flatten_dictionary - Функция, которая "выравнивает" словарь, делая его одноуровневым

    _onKeyRelease - функция, которая позволяет копировать, вырезать и вставлять текст в оконной форме даже с кириллической раскладкой

    xlsx_file_beautifulication - делает экселевский файл "красивым" - пытается подогнать ширину столбцов под текст и делает везде все нужные границы

    from_base10_to_baseXX - Перевод числа из десятеричной системы счисления в заданную.

    from_baseXX_to_base10 - Перевод числа из заданной системы счисления в десятеричную.

    find_all_systems - Определяет количество систем, записанных внутри одного обозначения.

    do_something_fun - создаёт анимированный градиентный экран

Классы:
    Blank - Общий класс для данных, которые извлекаются из бланка-заказа.

    DocxExpand - Дополнительные методы или функции, которые должны расширить функционал модуля docx.

    BaseXX - Класс, представляющий собой число в заданной системе счисления.

    Base12 - Класс чисел в 12-ичной системе счисления.
"""

from time import perf_counter, sleep
from os import remove, path, getcwd
from pathlib import Path
from re import findall, fullmatch, split, IGNORECASE
from zipfile import BadZipFile
from itertools import takewhile, dropwhile, pairwise
from io import BytesIO, StringIO
from math import floor, trunc, ceil, log10, sqrt
from xml.etree import ElementTree as ET
from sys import platform
from copy import deepcopy
from uuid import uuid3, NAMESPACE_DNS
from json import dump, dumps
from typing import Literal

from pandas import DataFrame, ExcelWriter
from dadata import Dadata
from PySimpleGUI import popup, popup_yes_no, popup_get_text, theme, LOOK_AND_FEEL_TABLE
from docx import Document
from docx.text.paragraph import Paragraph
from docx.document import Document as Doc
from docx.opc.exceptions import PackageNotFoundError
from docx2txt import process
from tabulate import tabulate
from openpyxl import load_workbook
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
import xlsxwriter

if platform == 'win32':
    from win32com import client as wc

DADATA_TOKEN = "f9607f7223c70da82d49c81434d14fa7b9ab635e"
DADATA_SECRET = "c8f87dab214c110e816b26801647a3912ef61762"
ALL_DIGITS = '0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZБГДЁЖИЙЛПФЦЧШЩЪЫЬЭЮЯÞΣΨΩ'
SUPPORTED_EXCTENTIONS_FOR_BLANK = ('.docx', '.doc', '.rtf', '.pdf', '.xml')

# ===========================================================================================================

popup_true_false = lambda question:popup_yes_no(question) == 'Yes'
mul_1000 = lambda x:str(floor(float(x) * 1000))
div_1000 = lambda x:str(round(float(x) / 1000))

class print_debug_mode_on:
    """Своеобразный класс-функция, задача которой выводить на экран информацию тогда и только тогда, когда того требуется. Ведёт себя абсолютно точно как обычный Принт (закидывайте в него аргументы, разделитье и знак окончания строки), но чтобы его включить, надо вызвать метод debug_mode_tumbler(). Ничего не возвращает, не создаёт объект, после принта сразу удаляет самого себя в момент инициализации, не пригоден для того, чтобы это было полноценным объектом, лол
    """
    debug_mode_on=True

    def __init__(self, *args, sep=' ', end='\n') -> None:
        self.args = args
        self.sep = sep
        self.end = end
        if self.debug_mode_on:
            print(*self.args, sep=self.sep, end=self.end)
        del self
        pass

    @classmethod
    def debug_mode_tumbler(self):
        self.debug_mode_on = not self.debug_mode_on

    @classmethod
    def check_debug_mode(self):
        print(self.debug_mode_on)

    def __str__(self) -> str:
        return self.sep.join(self.args) + self.end
    
    def __repr__(self) -> str:
        return 'print_debug_mode_on(' + ', '.join(f"'{arg}'" if isinstance(arg, str) else str(arg) for arg in self.args) + f", sep={self.sep}, end={self.end})"
    pass

def checking_arguments(*args):
    """Проверка типа и значения переменной. Ещё может вывести длину переменной, если подобное возможо для данного типа переменой.
    """

    for arg in args:
        print(type(arg), arg, sep=': ')
        try:
            print(len(arg))
        except:
            print('Данная переменная не имеет длинны!')

def file_names_csv_and_xlsx(file:str):
    """Создание двух переменных-названий файлов с расширениями csv и xlsx

    Args:
        file (str): Название файла. Лучше всего писать его без расширения.

    Returns:
        file_csv, file_xlsx: Кортеж, где первый элемент - название файла с расширением csv, а второй - xlsx
    """

    file_name = Path(file).stem
    return file_name + '.csv', file_name + '.xlsx'

def line_by_line(string:list, file:str, to_print=False):
    """Запись строки в csv-файл в его конец, переменная messaging должна отвечать за то, нужно ли выводить сообщение о записи

    Args:
        string (list): Строка в виде списка, которую надо записать
        file (str): Название файла
        to_print (bool, optional): Если нужно вывести сообщение о том, что было записано и в какой файл, написать True. По умолчанию тут стоит False.
    """

    DataFrame([string]).to_csv(file, mode='a', header=False, index=False)
    if to_print:
        print(f'В файл {file} было записано: {string}')

def decision(question:str, the_type:type, limit_do=0, limit_up=0, use_pysimplegui=False) -> bool|int|float:
    """Функция, которая возвращает значение в зависимости от вопроса. Может возвращать данные разных типов (лучше логическое или число, типа, для строки можно просто сделать инпут), завичит от того, что ввести. Может осуществляться либо через всплывающие окна, либо через консоль. Если нужно ввести число, то ОБЯЗАТЕЛЬНО надо не забыть про ограничения числа, что ли

    Args:
        question (str): Вопрос, на который надо дать ответ
        the_type (type): Тип переменной, которой надо вернуть
        limit_do (int, optional): Ограничение числового значения снизу. По умолчанию 0.
        limit_up (int, optional): Ограничение числового значения сверху. По умолчанию 0.
        use_pysimplegui (bool, optional): Нужно ли использовать оконную форму (с ней удобней, правда!). По умолчанию False.

    Returns:
        the_type: Нужная переменная нужного типа
    """

    if the_type == bool:
        if use_pysimplegui:
            answer = popup_true_false(question)
        else:
            answer = ''
            while type(answer) != bool:
                answer = input(question + '\n')
                if answer.upper() in ('ДА', 'YES', 'LF', 'НУЫ', 'Y', 'Д', 'TRUE'):
                    answer = True
                elif answer.upper() in ('НЕТ', 'NO', 'YTN', 'ТЩ', 'N', 'Н', 'FALSE'):
                    answer = False
                else:
                    print('Введите ещё раз.')
    else:
        answer = ''
        while not isinstance(answer, the_type):
            answer = popup_get_text(question) if use_pysimplegui else input(question + '\n')
            if answer is None:
                return None
            try:
                answer = the_type(answer)
            except ValueError:
                err_mess = 'Некорректно введённое значение.'
                popup(err_mess) if use_pysimplegui else print(err_mess)
            else:
                if not (limit_do <= answer <= limit_up):
                    err_mess = 'Значение за пределами диапазона.'
                    answer = ''
                    popup(err_mess) if use_pysimplegui else print(err_mess)
    return answer

def the_new_order(file, fpos, tpos):
    """Позаимствованный из интернета код, вот его описание: Takes a list of ints, and inserts the fpos (from position) int, to tpos (to position)
    
    Важное примечание: если нужно просто вытащить данные через с разных листов Экселя, пользуйтесь пандасом, у него для этого есть все инструменты!

    Args:
        file (_type_): Сюда, по идее, записывается название файла
        fpos (_type_): Сюда - исходная позиция, откуда двигается
        tpos (_type_): Сюда - то, куда двигать (кажется)
    """

    print('Начинается работа по перемещению листов в файле', file)
    wb = load_workbook(filename=file, data_only=True)
    shlist = wb.sheetnames  # get current order sheets in workbook
    lst = []
    lpos = (len(shlist) - 1) # last position
    if lpos >= fpos > tpos >= 0:  # move from a high to low position
        for x in range(lpos+1):
            if x == tpos:
                lst.append(fpos)
            elif tpos < x <= fpos:
                lst.append(x-1)
            else:
                lst.append(x)
    if lpos >= tpos > fpos >= 0:  # move from a low to high position
        for x in range(lpos+1):
            if x == tpos:
                lst.append(fpos)
            elif fpos <= x < tpos:
                lst.append(x+1)
            else:
                lst.append(x)
    wb._sheets = [wb._sheets[i] for i in lst]  # get each object instance from  wb._sheets, and replace
    wb.save(filename=file)
    print('Заканчивается работа по перемещению листов в файле', file)

def csv_cleaning(file_csv:str):
    """Функция очистки csv-файла - в целом, морально устарела

    Args:
        file_csv (str): название csv-файла
    """

    dataframe=open(file_csv, 'w+')
    dataframe.seek(0)
    dataframe.close

def ideal_message(curr_index:int, length:int, measure_genitive:str, time_start:float, show_left_time=True, return_message=False)-> (str | None):
    """Печатает строку состояния работы цикла в виде "Сделано столько-то чего-то там, осталось столько-то чего-то там. Процентов: проценты. Прошло времени - столько-то секунд. Осталось столько-то секунд". Желательно писать это в самом конце рабочего поля цикла.

    Args:
        curr_index (int): текущий номер итерации
        length (int): общее количество ожидаемых итераций
        measure_genitive (str): единица измерения в родительном падеже и множественном числе.
        time_start (float): время начала работы цикла. Задать ПЕРЕД всем циклом в формате "переменная = time.perf_counter()"
        show_left_time (bool, optional): На случай, если необходимо получить количество оставшегося времени. Эффективно применять, если массив большой, а время работы предполагается гигантским, на малых сроках он работает не особо хорошо. По умолчанию True.
        return_message (bool, optional): Надо ли возвращать сообщение как строку. По умолчанию False.

    Returns:
        (str | None): строка или вывод строки
    """
    
    all_time = round((perf_counter() - time_start), 2)
    procents = round((100 * curr_index / length), len(str(length)) - 2)

    ideal_message = f'Сделано {curr_index} {measure_genitive}, осталось {(length - curr_index)} {measure_genitive}. Процентов: {procents}. Прошло времени: {all_time} секунд.'
    if show_left_time:
        ideal_message += f' Осталось {round(all_time / (procents / 100) - all_time, 2)} секунд.'
    return ideal_message if return_message else print(ideal_message)

def dadata_inn_and_address(inn_or_address:str, only_main=True, is_it_inn=True, time_left=30):
    """Функция, обращающаяся к Дадате и забирающая оттуда всю необходимую информацию по данному ИНН

    Args:
        inn_or_address (str): Сюда нужно написать либо ИНН компании, либо адрес. ВАЖНО! Тип - str!
        only_main (bool, optional): Прописать False, если нужно забрать ещё и филиалы. По умолчанию True.
        is_it_inn (bool, optional): Прописать False, если нужно обратиться по адресу. По умолчанию True.
        time_left (int, optional): Время задержки в случае вылета. По умолчанию 30.

    Returns:
        result[0]['data']: словарь всех данных
        result: абсолютно все результаты
        None: если ничего не найдено
    """

    dadata = Dadata(DADATA_TOKEN)
    yes_result = False
    while yes_result == False: 
        try:
            result = dadata.find_by_id(name="party", query=inn_or_address) if is_it_inn else dadata.suggest("address", inn_or_address)
        except Exception as error:
            print(error)
            while time_left > 0:
                print('Осталось', time_left, 'сек')  
                sleep(1)
                time_left -= 1
            print('Новая попытка')
        else:
            yes_result = True
    
    if result != []:
        if only_main:
            return result[0]['data']
        else:
            return result
    else:
        return None

def dadata_left(to_print=False):
    """Возвращает количество оставшихся запросов

    Args:
        to_print (bool, optional): Надо ли сразу печатать оставшееся количество. По умолчанию False.

    Returns:
        int: left - количество оставшихся запросов
    """

    with Dadata(DADATA_TOKEN, DADATA_SECRET) as dadata:
        left = 100_000 - int(dadata.get_daily_stats()['services']['suggestions'])
        if to_print:
            print('Осталось запросов:', left)
        return left

def current_database():
    """Возвращает параметры базы данных в PostgreSQL

    Returns:
        connection: connect - Параметры базы данных, необходимо присвоить переменой, чтобы потом пользоваться этим соединением
    """

    return __import__('psycopg2').connect(
        user="postgres",
        password="p@ssw0rd",  # пароль, который указали при установке PostgreSQL
        host="192.168.30.223",
        port="5432",
        database="VEZA")

def veza_design():
    """Простенький метод, который задаёт созданное Майей оформление для программок
    """

    LOOK_AND_FEEL_TABLE['MyCreatedTheme'] = {
    'BACKGROUND': '#BECBBA',
    'TEXT': '#172412',
    'INPUT': '#ECF4E7',
    'TEXT_INPUT': '#172412',
    'SCROLL': '#172412',
    'BUTTON': ('#172412', '#EEFFFF'),
    'PROGRESS': ('#172412', '#EEFFFF'),
    'BORDER': 3, 
    'SLIDER_DEPTH': 2, 
    'PROGRESS_DEPTH': 2, }  
    # Switch to use your newly created theme
    theme('MyCreatedTheme')

def convert_using_win32(old_filename:str|Path, new_filename:str|Path, extension:str, to_print=True):
    """Конвертация всего, что можно открыть в Ворде, в docx. Иногда оно можно не сработать по непонятным мне причинам - я попытался это предусмотреть, однако есть ещё вероятность, что что-то может пойти наперекосяк. Модуль win32 - мощный инструмент. Возможно, даже слишком мощный. Не гарантируется работоспособность на других операционных системах

    Args:
        old_filename (str|Path): Имя старого файла
        new_filename (str|Path): Имя нового файла
        extension (str): в какое расширение нужно сконвертировать. Поддерживаемые варианты: 'docx', 'xlsx'
    """
    # old_filename, new_filename = Path(old_filename), Path(new_filename)
    print_debug_mode_on(old_filename, new_filename, sep='\n')

    path_for_wc = lambda x: str(x.absolute()).replace('/', '\\') if isinstance(x, Path) else x.replace('/', '\\')

    if to_print:
        print('Начинается конвертация файла', old_filename)
    match extension:
        case '.docx':
            app = 'Word.Application'
            save_code = 16
        case '.xlsx':
            app = 'Excel.Application'
            save_code = 51
    print_debug_mode_on(old_filename.absolute())
    doc = wc.Dispatch(app).Documents.Open(path_for_wc(old_filename))
    if doc is not None:
        doc.SaveAs(path_for_wc(new_filename), save_code)
        doc.Close(False)
        if to_print:
            print('Конвертирован файл', old_filename.__str__())
    else:
        print('Не удалось сконвертировать файл! Попробуйте сделать вручную!')

def ordered_content_from_docx(value:str|Path|Doc, save_as_objects=False) -> tuple:
    """Функция читает подряд docx-файл и возвращает все основные объекты файла в порядке их появления. Это серьёзно доработанный код из интернета, из которого удалено всё лишнее и, на мой скромный вкус, не значащее. Исходный код лежит в соседнем файле. Главное улучшение - уход от датафреймовой структуры в пользу более примитивной, но более надёжной структуры двухмерных списков-массивов, что, на мой скромный вкус, лучше. Позволяет, как минимум, уйти от необходимости импортировать функцию по игнорированию ФьючерВорнинг.

    Args:
        value (str|Path|Doc): вордовский файл или объект
        save_as_objects (bool): надо ли сохранять параметры в список как объекты модуля докс или как текст или ещё что? По умолчанию False

    Raises:
        ValueError: _description_

    Returns:
        tuple: document, combined_list, image_list: объект класса Document для дальнейшей с ним работы и два списка - общий и картиночек

    Yields:
        _type_: каждый параграф и таблицу в порядке появления в документе
    """
    print_debug_mode_on(value)

    if isinstance(value, (str, Path)):
        document = Document(value)
    elif isinstance(value, Doc):
        document = value
    else:
        raise TypeError("Неподдерживаемый тип")

    ##This function extracts the tables and paragraphs from the document object
    def iter_block_items(parent):
        """Yield each paragraph and table child within *parent*, in document order. Each returned value is an instance of either Table or Paragraph. *parent* would most commonly be a reference to a main Document object, but also works for a _Cell object, which itself can contain paragraphs and tables.

        Args:
            parent (_type_): _description_

        Raises:
            ValueError: _description_

        Yields:
            Paragraph | Table: параграф-объект или таблица-объект
        """
        if isinstance(parent, Doc):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    combined_list = []
    image_list = []
    i = 0
    imagecounter = 0

    for block in iter_block_items(document):
        if 'text' in str(block):
            isappend = False
            
            runboldtext = ''
            for run in block.runs:                        
                if run.bold:
                    runboldtext = runboldtext + run.text
                    
            style = str(block.style.name)

            if save_as_objects:
                appendtxt = block
            else:
                appendtxt = str(block.text)
                appendtxt = appendtxt.replace("\n","")
                appendtxt = appendtxt.replace("\r","")
            tabid = 'Novalue'
            
            isappend = True
            for run in block.runs:
                xmlstr = str(run.element.xml)
                my_namespaces = dict([node for _, node in ET.iterparse(StringIO(xmlstr), events=['start-ns'])])
                if 'pic:pic' in xmlstr:
                    for pic in ET.fromstring(xmlstr) .findall('.//pic:pic', my_namespaces):
                        name_attr = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces).get("name")
                        embed_attr = pic.find("pic:blipFill/a:blip", my_namespaces).get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        isappend = True
                        appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                        image_list.append([
                            imagecounter,
                            embed_attr,
                            name_attr,
                            __import__('base64').b64encode(document.part.related_parts[embed_attr]._blob).decode()])
                        style = 'Novalue'
                    imagecounter = imagecounter + 1
                
        elif 'table' in str(block):
            isappend = True
            style = 'Novalue'
            appendtxt = block
            tabid = i
            i += 1
        if isappend:
            combined_list.append([appendtxt, tabid, style])

    return document, combined_list, image_list

def flatten_dictionary(dictionary:dict, use_new_key:bool, parent_key='', sep='_') -> dict:
    """Функция, которая "выравнивает" словарь, делая его одноуровневым. Возможно либо сохранение внутренних названий ключей, либо же полное наименование ключей. Старый словарь не изменяется, создаётся новый словарь

    Args:
        dictionary (dict): словарь, который надо "выравнять"
        use_new_key (bool): надо ли использовать новые ключи
        parent_key (str, optional): "родительский" ключ. По умолчанию ''.
        sep (str, optional): Разделитель. По умолчанию '_'.

    Returns:
        dict: выровненный словарь
    """
    items = []
    for key, value in dictionary.items():
        new_key = parent_key + sep + key if (parent_key and use_new_key) else key
        items.extend(flatten_dictionary(value, use_new_key, new_key, sep=sep).items()) if isinstance(value, dict) else items.append((new_key, value))
    return dict(items)

def _onKeyRelease(event):
    """Взятый из интернета код, который позволяет копировать, вырезать и вставлять текст в оконной форме даже с кириллической раскладкой

    Args:
        event (_type_): _description_
    """
    ctrl  = (event.state & 0x4) != 0
    if event.keycode==88 and ctrl and event.keysym.lower() != "x":
        event.widget.event_generate("<<Cut>>")
    if event.keycode==86 and ctrl and event.keysym.lower() != "v":
        event.widget.event_generate("<<Paste>>")
    if event.keycode==67 and ctrl and event.keysym.lower() != "c":
        event.widget.event_generate("<<Copy>>")

def progress_bar_updater(progress_bar, c:int, negative=False):
    """Просто обновление счётчика. Увеличивает (и ли уменьшает) его его на один и обновляет прогресс-бар. Вынесено в модуль, потому что слишком часто приходится проделывать эту процедуру в моих программах...
        
    Returns:
        int: c - новое значение счётчика, увеличенное на единицу

    Args:
        progress_bar (_type_): прогрессбар
        c (int): текущее значение счётчика
        negative (bool, optional): если надо, чтобы счётчик уменьшался, параметр должен быть равен True. По умолчанию равен False.

    Returns:
        tuple[Any, int]: progress_bar - изменённый прогрессбар, переданный функции; c - изменённый счётчик
    """
    c = c - 1 if negative else c + 1
    progress_bar.UpdateBar(c)
    return progress_bar, c

def xlsx_file_beautifulication(file_name:str|BytesIO, df_name) -> (BytesIO | None):
    """Функция делает экселевский файл "красивым" - пытается подогнать ширину столбцов под текст и делает везде все нужные границы.

    Args:
        file_name (str|BytesIO): имя изменяемого файла или файло-объект
        df_name (_type_): передаваемый датафрейм, который надо внести

    Returns:
        BytesIO | None: файло-объект, если он был передан - в ином случае, ничего не возвращаем
    """
    no_permission_error = False  # И мы тут снова задействуем ту переменную и ту логику, что будем пытаться получить доступ к файлу на случай, если я его не закрыл
    while not no_permission_error:
        try:
            writer = ExcelWriter(file_name, engine='xlsxwriter')  # Дальше идут довольно абстрактные строчки, потому что я их взял из интернета. Вот вы знали, что в пандасе есть ЭксельРайтер? теперь знаете. И вот я не знал
        except PermissionError:  # Вообще сомневаюсь, что кто-то будет держать сводную таблицу открытой, но во избежание это надо сделать, потому что обычно сводную таблицу забываю закрыть я
            popup('Закройте файл!')  # Просто всплывающее окошко - назойливое, но настойчивое
        else:
            no_permission_error = True
    no_permission_error = False  # Ещё раз перезадаю переменную, а то мало ли

    df_name.to_excel(writer, sheet_name='sheetName', index=False, na_rep='NaN')  # Ну, тут я верю интернету наслово. Если оно и в самом деле сделает то, что надо, то ладно. Тут мы, типа, передаём значения в сам эксель (мне казалось, это задаётся немного иначе, ну ладно)
    worksheet = writer.sheets['sheetName']
    for column in df_name:  # А дальше логика такая - находим максимальную ширину, которая либо в текущей строке, либо в шапке. и перезадаём ширину всей таблички. Я верю интернету наслово - оно реально работает
        column_length = max(df_name[column].astype(str).map(len).max(), len(column)) + 3
        col_idx = df_name.columns.get_loc(column)
        worksheet.set_column(col_idx, col_idx, column_length)
    # workbook = writer.book
    # border_fmt = workbook.add_format({'bottom':1, 'top':1, 'left':1, 'right':1})
    # worksheet.conditional_format(xlsxwriter.utility.xl_range(0, 0, len(session_state['pivot_table']), len(session_state['pivot_table'].columns)), {'type': 'no_errors', 'format': border_fmt})
    worksheet.conditional_format(xlsxwriter.utility.xl_range(0, 0, len(df_name), len(df_name.columns)), {'type': 'no_errors', 'format': writer.book.add_format({'bottom':1, 'top':1, 'left':1, 'right':1})})
    writer.close()  # В старой версии Пандаса тут надо было писать сейв, а в новой - клоуз
    return file_name if isinstance(file_name, BytesIO) else None

def find_all_systems(main_system_name:str) -> list:
    """Определяет количество систем, записанных внутри одного обозначения. Так как обозначения могут быть почти что любыми, функция будет постоянно дорабатываться

    Args:
        main_system_name (str): Основное обозначение системы, из которого будут извлекаться все прочие

    Returns:
        list: all_system_names - все названия систем
    """

    system_names = main_system_name.split(',')  # Самая важная часть - разные системы могут быть перечислены через запятую. Но это не всегда так - чтобы не писать подряд идущие номера систем, их пишут через дефис (зачастую без пробелов). Так что, нужно дальше сплитать по дефису? Ни-хе-ра - в самом номере системы может содержаться дефис, и тогда что прикажете делать?
    print_debug_mode_on(system_names)
    all_system_names = []  # На последний вопрос есть ответ, но пока нам нужен перечень ВСЕХ систем, так что создаём пустой список
    for system_name in system_names:  # Пробегаемся по всему, что разделено запятыми
        system_name = system_name.strip()
        is_match = fullmatch(r'(([A-Za-zА-Яа-яЁё])\d{1,})\s*?-\s*?(\2\d{1,})', system_name)  # А теперь кое-что сложное. Номера систем могут быть записаны как угодно, одному богу известно, чем руководствуются люди, записывая их. И я научился отслеживать все более-менее сложные случаи. Но самый банальный, когда у нас записано, например, "П1-П4", программа не отслеживала. Поэтому я допёр до использования дополнительной функции из ре-модуля, которая проверяет на соответствие описанной схеме
        if is_match:  # И если да...
            system_name_all = is_match.group(1, 3)  # Забираем первую и третью группы - ну, так работает этот модуль
        else:  # В ином случае...
            is_match = fullmatch(r'(.+?)-(\d{1,})-(\d{1,}).*', system_name)
            if is_match:
                print_debug_mode_on(is_match.group(1, 2, 3))
                return [f"{is_match.group(1)}-{i + 1}" for i in range(int(is_match.group(2)), 1 + int(is_match.group(3)))]
            all_masks = (r'\d?[A-Za-zА-Яа-яЁё]\D?\D?\D?\D?\S?\S?\S?\S?\d{1,}[А-Яа-яЁё]*', r'\d{9}[а-яё]?')  # У нас есть набор масок, и нам их надо будет перебрать. Я буду честен - я сам не понимаю, как работает первое регулярное выражение, но оно просто работает и почти безошибочно находит все номера, причём умея отличать, где заканчивается номер в тех случаях, когда и в номере дефис, и между номерами дефис. Полагаю, что если что-то пойдёт не так и не найдётся, надо будет пошаманить ещё с выражением, но пока в этом нет нужды
            system_name_all = max((findall(the_mask, system_name, IGNORECASE) for the_mask in all_masks), key=lambda x:len(x))  # И нам нужно найти то, в котором будет найден максимум. Раньше, справедливости ради, тут было три варианта масок, но одна из них толком не работала, поэтому от неё отказался.
        system_name_all = system_name_all if system_name_all else [main_system_name]  # Иногда бывает по-идиотски записанная система, где нет цифр в конце, так что мы просто считаем, что в таком случае всего одна система и успокаиваемся
        print_debug_mode_on(system_name_all)
        if len(system_name_all) > 1:  # А вот если у нас нашлось больше одной, это значит, что у нас записаные через дефис системы. которых дофига. И тут могут возникнуть ещё сложности. Главная сложность - как оказалось, через дефис может быть не только формата "П1.Х-П1.У", но и формата "ПХ.1-ПУ.1". Вообще, как я понял, сильнее всего отличаться будет некое число, которое и обозначает начало и конец последовательности. Поэтому нам надо вообще понять, что с этими числами не так по ходу дела
            all_positions_in_system_name = tuple(tuple(filter(None, findall(r'(\D?|\d+)', sys_name, IGNORECASE))) for sys_name in system_name_all)  # Итак, логика в чём? У нас есть два класса элементов - не-числа и числа. Не-числа всегда одинаковые. Числа могут и отличаться. Нам важно понять, что есть не-число, а что - число, и мы разбиваем всю строку на такие вот элементы
            print_debug_mode_on(all_positions_in_system_name)
            system_condition = lambda x: x[0] == x[1]  # Потом нашей задачей будет сравнить каждый из элементов начала и конца. Сравнивать будем по этой мини-функции. Логика в том, что отличаться должно некое число, но оно может быть как в конце, так и в середине. И поэтому нам понадобится три разных переменных, из которых средняя будет разбита на две переменные для удобства. Итак
            before_changing_part = ''.join(el[0] for el in takewhile(system_condition, zip(all_positions_in_system_name[0], all_positions_in_system_name[1])))  # Тейквайл берёт последовательность до момента, когда нарушается условие. Условие - одинаковость элементов начального и конечного обозначения, которые мы зазипали. И чтобы два раза не вставать ещё и объединим в строчку
            after_changing_part = dropwhile(system_condition, zip(all_positions_in_system_name[0], all_positions_in_system_name[1]))  # Дропвайл работает диаметрально противоположно тейквайлу - забирает все значения С МОМЕНТА, когда нарушается условие. В том числе и то место, где условие нарушается. Так что я решил воспользоваться этим. Мы создаём генератор, в котором нам нужен будет первый элемент, а остальные надо будет загнать в строку. Так что...
            changing_part_start, changing_part_thend = next(after_changing_part)  # Мы для обозначения начала и конца сдвигаем генератор, распаковывая кортеж (а первый элемент в генераторе у нас кортеж и есть). Раньше для задания этой функции использовался Фильтерфолс
            after_changing_part = ''.join(el[0] for el in after_changing_part)  # А оставшееся соединяем. Генератор исчерпан, радуемся жизни и восхищаемся великолепием данной оптимизации
            print_debug_mode_on(before_changing_part, changing_part_start, changing_part_thend, after_changing_part)
            print_debug_mode_on(changing_part_start.isdigit(), changing_part_thend.isdigit())
            if changing_part_start.isdigit() and changing_part_thend.isdigit():  # Внезапно всплыла проблема, что порой название системы может быть написано через пробел. Не беда - собственно говоря, вся эта херобористика заточена на то, что у нас в изменяемой части будет два числа. Так что логично, что мой великолепный алгоритм будет работать тогда и только тогда, когда оба этих параметра - числа
                all_system_names += [f"{before_changing_part}{'0' * (min(len(changing_part_start), len(changing_part_thend)) - len(str(i)))}{i}{after_changing_part}" for i in range(int(changing_part_start), int(changing_part_thend) + 1)]  # А потом мы создаём все промежуточные обозначения систем. Формула выглядит очень мудрёной, но, думаю, в ней можно разобраться: соединяем начало, сколько-то нулей (пока нулей, но допускаю, что могут быть и иные случаи), число и конец.
            else:  # А если это у нас не числа, то уходим отсюда, нам тут делать нечего
                all_system_names = system_names
                break
            print_debug_mode_on(all_system_names)
        else:
            all_system_names += system_name_all  # Ну а если у нас одна система, то её и добавляем
    return all_system_names

def document_with_exception(file, to_send_message):
    try:
        return Document(file)
    except (PackageNotFoundError, ValueError, BadZipFile):
        to_send_message(Exception(f'Преобразуйте файл {file} в docx-формат, после чего попробуйте ещё раз!'))
    except Exception as err:
        to_send_message(err)

def round_up(n, decimals = 0):
    """Не совсем понятно зачем написаная Саакаяном функция

    Args:
        n (_type_): _description_
        decimals (int, optional): _description_. Defaults to 0.

    Returns:
        _type_: _description_
    """
    multiplier = 10 ** decimals 
    return ceil(n * multiplier) / multiplier

def same_latin_cyril(string:str, from_latin_to_cyril=True) -> str:
    """Смена всех похожих друг на друга символом латиницы и кириллицы друг на друга

    Args:
        string (str): строка, которую надо изменить
        from_latin_to_cyril (bool, optional): По умолчанию True, и функция меняет латиницу на кириллицу. Если будет False, то будет наоборот.

    Returns:
        str: string, но изменённая
    """

    for latin, cyril in zip(('A', 'B', 'E', 'K', 'M', 'H', 'O', 'P', 'C', 'T', 'Y', 'X'), ('А', 'В', 'Е', 'К', 'М', 'Н', 'О', 'Р', 'С', 'Т', 'У', 'Х')):
        string = string.replace(latin, cyril) if from_latin_to_cyril else string.replace(cyril, latin)
    return string

def poles_finder(ndv:int|float) -> Literal[2, 4, 6, 8]:
    """ Определяем количество полюсов - оно зависит от количества оборотов в минуту

    Args:
        ndv (int | float): обороты в минуту

    Returns:
        Literal[2, 4, 6, 8]: количество полюсов
    """

    if 1500 < ndv:
        return 2
    elif 1000 < ndv <= 1500:
        return 4
    elif 750 < ndv <= 1000:
        return 6
    else:
        return 8

class Blank:
    """Общий класс для данных, которые извлекаются из бланка-заказа. В настоящий момент поддерживает бланки на ВЕРОСА, канальное оборудование, осевые вентиляторы, индустриальные вентиляторы. Чем новее бланк, тем лучше. Бланк ОБЯЗАТЕЛЬНО должен быть в форматe docx или doc!!! При попытке загнать файл другого формата будет ошибка, так что будьте внимательны. Если же файл формата doc, то могут возникнуть непредвиденные ошибки, будьте осторожны!
    В качестве аргумента принимается имя файла с бланком.
    """
    ready_blank_types = ('Другое', 'ВЕРОСА', 'Общепромышленные', 'Канальное оборудование', 'Индустриальный вентилятор', 'Индивидуальный тепловой пункт', 'Холодильное оборудование', 'Драйкулеры')
    all_columns = (
            'Бланк-заказ', 'Дата бланк-заказа', 'Входящий номер', 'Дата входящего номера', 'Объект', 'Номер объекта', 'Дата', 'Организация', 'Менеджер', 'Выполнил', 'Поток', 'Название', 'Типоразмер',  # Колонки основной информации
            'Назначение', 'Название блока', 'Тип блока', 'Информация о блоке'  # Колонки информации с содержимым бланка
        )

    def __init__(self, inputed:str|Path|Doc|BytesIO) -> None:
        """Инициализация бланка

        Args:
            inputed (str | Doc): переданный параметр - либо адрес файла, либо документ-объект

        Raises:
            TypeError: возникает, если передано что-то не строковое и не объектно-документное, либо если файл не того формата. Возникновени других ошибок означает непредусмотренные ошибки
        """

        if isinstance(inputed, (str, Path)):
            self.filename = Path(inputed)
            if self.filename.suffix.lower() in SUPPORTED_EXCTENTIONS_FOR_BLANK[0:4]:
                if self.filename.suffix.lower() == '.docx':
                    self.Document_object = Document(self.filename)
                    self.temporary_filename = self.filename
                else:
                    try:
                        self.Document_object = Document(self.filename)
                    except (PackageNotFoundError, ValueError, BadZipFile):
                        self.temporary_filename = Path(getcwd() + '\\' + str(self.filename.name)).with_suffix('.docx')
                        convert_using_win32(self.filename, self.temporary_filename, '.docx', False)
                        self.Document_object = Document(self.temporary_filename)
                    else:
                        self.temporary_filename = self.filename
            else:
                raise TypeError("Введённый файл не является файлом поддерживаемых расширений (" + ','.join(SUPPORTED_EXCTENTIONS_FOR_BLANK[0:3]) + ")!") if self.filename.suffix.lower() != '.xml' else TypeError("Для обработки xml-файлов используйте класс XmlParer или производные от него!")
        elif isinstance(inputed, Doc):
            self.filename, self.temporary_filename = None, None
            self.Document_object = inputed
        elif isinstance(inputed, BytesIO):
            self.Document_object = Document(inputed)
            self.filename, self.temporary_filename = (inputed.name, ) * 2 if hasattr(inputed, 'name') else (None, None)
        else:
            raise TypeError("Переданный аргумент некорректного типа!")

        bio = BytesIO()
        self.Document_object.save(bio)
        # self.qweqwe = process(bio).split('\n')
        self.docx_text = [info.strip() for info in process(bio).split('\n') if info and not set(info).issubset({' ', '\xa0', '\t'})]

        self.blank_type = {key: False for key in self.ready_blank_types}
        self.main_information, self.all_avaiable_information = self.__blank_processing()
        self.IS_OTHERS, self.IS_VEROSA, self.IS_OBPROM, self.IS_CHANAL, self.IS_INDUST, self.IS_INTEPU, self.IS_KHOLOD, self.IS_DRYDOL = self.blank_type.values()
        self.ALL_MAIN_INFO = {key: value for key, value in zip((data['Название блока'] for data in self.all_avaiable_information), (data['Информация о блоке'] for data in self.all_avaiable_information))}
        self.all_main_info_text = tabulate(self.ALL_MAIN_INFO, ('Заголовок', 'Значение'))

    def __str__(self) -> str:
        return next(key for key, value in self.blank_type.items() if value) + '\n\n' + tabulate(self.main_information.items(), ('Заголовок', 'Значение'))

    def __blank_processing(self):
        """Основная функция обработки бланков. Склеена из кусков разной степени древности, где-то оно модернизировано, где-то - нет. Но оно работает практически стабильно. Разбирает бланк почти по винтику и вытаскивает всю основную информацию (название бланк-заказа, дата, входящий номер, дата входящего номера, объект, номер объекта, дата, организация, менеджер, испиолнитель, название, типоразмер) и все параметры установки/оборудования/вентилятора.

        Returns:
            dict[str, str], list: _description_
        """

        def find_system_name() -> str:
            """Вспомогательная функция, которая помогает задать название системы, если оно оказалось не задано. Обычно это в бланках канального или общепрома. Тут решение простое - система записана немного в другом месте бланка, не так, где этого ищу я

            Returns:
                str: Найденное название системы (или исходное, если не было найдено или не надо было вовсе искать)
            """
            print_debug_mode_on(result['Название'])
            if result['Название'] == '-':
                elem = {findall(r'БЛАНК-ЗАКАЗ\s?(.+?)\s?от', cell_text, IGNORECASE)[0] for cell_text in self.docx_text if 'БЛАНК-ЗАКАЗ' in cell_text.upper()}
                print_debug_mode_on(elem)
                return tuple(elem)[0] if elem else result['Название']  # Если что-то нашлось, сохраняем как кортеж (у нас на элем будет либо множеством, либо списком, так что сработает) и забираем первый элемент (обычно он же единственный)
            else:
                return result['Название']

        count_kanal = 0
        print_debug_mode_on(self.docx_text)
        for cell_text in self.docx_text:
            if any(self.blank_type.values()):
                break
            if 'Кондиционеры центральные каркасно-панельные' in cell_text or 'Кондиционеры компактные панельные' in cell_text:
                self.blank_type[self.ready_blank_types[1]] = True
            
            if 'ОБЩЕПРОМЫШЛЕННЫЕ И СПЕЦИАЛЬНЫЕ ВЕНТИЛЯТОРЫ ВЕЗА' in cell_text:
                self.blank_type[self.ready_blank_types[2]] = True

            if cell_text in ['ООО «ВЕЗА»', '111397, Москва, Зеленый пр-т, д20, 6 этаж', 'Тел: +7(495)989-47-20; Факс: +7(495)626-99-02', 'veza@veza.ru'] or cell_text in ['ООО "Веза"', 'Москва, Зеленый проспект д.20', 'Тел: +7 (495) 989-47-20; Факс: +7 (495) 989-47-20', 'msk1@veza.ru']:
                count_kanal += 1
            if count_kanal >= 4:
                self.blank_type[self.ready_blank_types[3]] = True

            if cell_text == 'Технические характеристики на стандартный индустриальный вентилятор':
                self.blank_type[self.ready_blank_types[4]] = True

            if 'Пункт тепловой индивидуальный' in cell_text:
                self.blank_type[self.ready_blank_types[5]] = True

            if 'Чиллер серии' in cell_text:
                self.blank_type[self.ready_blank_types[6]] = True

            if 'драйкулер' in cell_text:
                self.blank_type[self.ready_blank_types[7]] = True
            pass
        self.blank_type[self.ready_blank_types[0]] = not any(self.blank_type.values())
        print_debug_mode_on(self.blank_type)
        pass

        result = {key: '-' for key in self.all_columns[0:-4]}
        ALL_RESULTS = []
        clean_information = {key: '-' for key in self.all_columns[-4:]}

        if not self.blank_type[self.ready_blank_types[0]]:
            to_create_table = lambda start, end: tuple(takewhile(lambda x: end not in x, dropwhile(lambda x: start not in x, self.docx_text)))

            if self.blank_type[self.ready_blank_types[1]] or self.blank_type[self.ready_blank_types[2]]:
                for cell_text in self.docx_text:
                    if all(value != '-' for value in result.values()):
                        break

                    if 'БЛАНК' in cell_text.upper():
                        result['Бланк-заказ'], result['Дата бланк-заказа'] = findall(r'БЛАНК[\-\s]ЗАКАЗ\s?(.+?)\s?от\s?(\d{2}\.\d{2}\.\d{4})', cell_text, IGNORECASE)[0]
                    if 'входящий:' in cell_text.lower():
                        result['Входящий номер'], result['Дата входящего номера'] = findall(r'входящий:\s?(.+?)\s?от\s?(\d{2}\.\d{2}\.\d{4})', cell_text, IGNORECASE)[0]
                    if 'ОБЪЕКТ:' in cell_text.upper():
                        result['Объект'] = findall(r'объект:\s?(.+)', cell_text, IGNORECASE)[0]
                        if '(' in cell_text:
                            result['Номер объекта'] = findall(r'\((.+?)\)', cell_text, IGNORECASE)
                            if result['Номер объекта']:
                                result['Номер объекта'] = result['Номер объекта'][-1]
                            else:
                                result['Номер объекта'] = findall(r'\((.+)', cell_text, IGNORECASE)[-1]
                    if 'код:' in cell_text.lower():
                        result['Типоразмер'] = findall(r'код:\s?(.+)', cell_text, IGNORECASE)[0]

                    for column_key in self.all_columns[6:-3]:
                        if column_key.lower() + ':' in cell_text.lower():
                            print_debug_mode_on(self.all_columns, self.all_columns[6:-3], cell_text)
                            result[column_key] = findall(fr"{column_key}:\s?(.+)", cell_text, IGNORECASE)
                            result[column_key] = result[column_key][0] if result[column_key] else '-'
                    pass

                for result_key in result:
                    # print_debug_mode_on(result_key, result_file[result_key])
                    abzatc = result[result_key].split('\n')
                    print_debug_mode_on(abzatc)
                    if len(abzatc) > 1:
                        if result_key.upper() in result[result_key].upper():
                            for strochka in abzatc:
                                print_debug_mode_on(strochka)
                                if result_key.upper() in strochka.upper():
                                    result[result_key] = strochka.split(':')[1].lstrip()
                        else:
                            result[result_key] = abzatc[0].lstrip()

                result['Название'] = find_system_name()

                print_debug_mode_on(result)
                pass

                if self.blank_type[self.ready_blank_types[1]]:

                    monoblocks, blocks = 0, 0
                    # mono_flag, bloc_flag = False, False
                    slash_checked = False
                    for cell_text in self.docx_text:
                        print_debug_mode_on(cell_text)
                        pass
                        if monoblocks and blocks:
                            break

                        has_bloks = 'блоков' in cell_text.lower() and cell_text != 'Наименование блоков с индексами и характеристиками входящего оборудования' and 'моноблоков' not in cell_text.lower()
                        has_monos = 'моноблоков' in cell_text.lower()
                        print_debug_mode_on(has_bloks, has_monos)

                        if has_bloks and has_monos:
                            res = findall(r'блоков=(\d+).+моноблоков=(\d+)', cell_text, IGNORECASE)
                            print_debug_mode_on(res)
                            if res:
                                blocks, monoblocks = map(int, res[0])
                                print_debug_mode_on(blocks, monoblocks)
                                pass
                            else:
                                pass
                        elif has_monos:
                            monoblocks = int(findall(r'моноблоков=(\d+)шт', cell_text, IGNORECASE)[0])
                            print_debug_mode_on(monoblocks)
                            pass
                        elif has_bloks:
                            blocks = int(findall(r'блоков=(\d+)шт', cell_text, IGNORECASE)[0])
                            print_debug_mode_on(blocks)
                            pass
                        else:
                            pass

                    print_debug_mode_on(monoblocks, blocks)
                    pass

                    all_needed_information = []
                    for i in range(1, monoblocks + 1):
                        dob_mono = str(i) + '.'
                        needed_information = {
                            'value' : dob_mono,
                            'type' : 'моноблок',
                            'title' : '',
                            'found_value' : False,
                            'information' : [],
                            'found_info' : False}
                        all_needed_information.append(needed_information)
                        for j in range(1, blocks + 1 + 1):
                            dob_bloc = dob_mono + str(j) + '.'
                            needed_information = {
                                'value' : dob_bloc,
                                'type' : 'блок',
                                'title' : '',
                                'found_value' : False,
                                'information' : [],
                                'found_info' : False}
                            all_needed_information.append(needed_information)
                    # print_debug_mode_on(all_needed_information)
                    all_headers = [need_info['value'] for need_info in all_needed_information]
                    print_debug_mode_on(all_headers, *self.docx_text, sep='\n||')
                    pass

                    for i in range(len(all_needed_information) - 1):
                        header = all_needed_information[i]['value']
                        head = header.replace('.', '\.')
                        print_debug_mode_on(header, head)
                        work_zone = list(dropwhile(lambda x: not findall(fr"^{head}\s.+", x), self.docx_text))
                        print_debug_mode_on(work_zone)
                        if work_zone:
                            for ender in all_headers[all_headers.index(header) + 1:]:
                                end = ender.replace('.', '\.')
                                print_debug_mode_on(ender, end)
                                new_word_zone = list(takewhile(lambda x: not findall(fr"^{end}\s.+", x), work_zone))
                                print_debug_mode_on(new_word_zone)
                                if new_word_zone != work_zone:
                                    all_needed_information[i]['title'] = new_word_zone[0]
                                    all_needed_information[i]['found_value'] = True
                                    all_needed_information[i]['information'] = '; '.join(new_word_zone[1:])
                                    print_debug_mode_on(all_needed_information[i])
                                    all_needed_information[i]['found_info'] = True
                                    break
                                else:
                                    pass
                                pass
                            else:
                                possible_finish_criterion = ('Автоматика', 'Должность,ФИО,подпись')
                                for pfc in possible_finish_criterion:
                                    new_word_zone = list(takewhile(lambda x: pfc not in x, work_zone))
                                    print_debug_mode_on(new_word_zone)
                                    if new_word_zone != work_zone:
                                        all_needed_information[i]['title'] = new_word_zone[0]
                                        all_needed_information[i]['found_value'] = True
                                        all_needed_information[i]['information'] = '; '.join(new_word_zone[1:])
                                        print_debug_mode_on(all_needed_information[i])
                                        all_needed_information[i]['found_info'] = True
                                        break
                                    else:
                                        print_debug_mode_on(new_word_zone)
                                        pass                                
                        pass
                    print_debug_mode_on(*all_needed_information, sep='\n')
                    pass

                    for additional_information in all_needed_information:
                        if (additional_information['found_value'] and additional_information['found_info']) or (additional_information['title'] and additional_information['information']):
                            clean_information = {
                                'Название блока' : additional_information['title'],
                                'Тип блока' : additional_information['type'],
                                'Информация о блоке' : additional_information['information']}
                            print_debug_mode_on(clean_information)
                            itog_information = result | clean_information
                            ALL_RESULTS.append(itog_information)
                else:
                    all_information = list(dropwhile(lambda x: '1. ' not in x, self.docx_text))
                    print_debug_mode_on(all_information)

                    thend_criterion = 'Дополнительное оборудование'

                    predicate_for_end = lambda x: x != 'Дополнительное оборудование'
                    predicate_for_sta = lambda x: x != thend_criterion

                    clean_information = {
                        'Название блока' : all_information[0],
                        'Тип блока' : 'моноблок',
                        'Информация о блоке' : '; '.join(el for el in takewhile(predicate_for_sta, all_information[1:]) if not el.isdigit() and all(el != e for e in ('Спектральные уровни звуковой мощности', 'Среднегеометрические частоты октавных полос, Гц', 'на входе, дБ', 'на выходе, дБ')))}
                    print_debug_mode_on(clean_information)
                    itog_information = result | clean_information
                    ALL_RESULTS.append(itog_information)

                    add_info = tuple(item for item in all_information if 'ополнительн' in item)
                    print_debug_mode_on(add_info)
                    if add_info:
                        thend_criterion = findall(r'стр 1 \/ \d{1,}', '; '.join(all_information))[0]
                        print_debug_mode_on(thend_criterion)
                        pass

                        clean_information = {
                            'Название блока' : 'Дополнительное оборудование',
                            'Тип блока' : 'моноблок',
                            'Информация о блоке' : '; '.join(takewhile(predicate_for_sta, tuple(dropwhile(predicate_for_end, all_information))[1:]))}
                        print_debug_mode_on(clean_information)
                        itog_information = result | clean_information
                        ALL_RESULTS.append(itog_information)
                    pass
                pass

            elif self.blank_type[self.ready_blank_types[3]]:
                indexes = []
                print_debug_mode_on(self.docx_text)
                for first_cell, next_cell in pairwise(self.docx_text):
                    print_debug_mode_on(first_cell, next_cell)
                    prov = first_cell.upper()

                    if 'ПРОЕКТ' in prov:
                        result['Входящий номер'] = cell_text.split()[1]

                    if first_cell in ['Объект:', 'Заказчик:', 'Исполнитель:', 'Название:']:
                        match first_cell:
                            case 'Объект:':
                                result['Объект'] = next_cell
                            case 'Заказчик:':
                                result['Организация'] = next_cell
                            case 'Исполнитель:':
                                result['Выполнил'] = next_cell
                            case 'Название:':
                                result['Бланк-заказ'] = next_cell
                                result['Название'] = next_cell

                    if 'Индекс:' in first_cell:
                        print_debug_mode_on(first_cell)
                        if first_cell.split()[1] not in indexes:
                            indexes.append(first_cell.split()[1])
                    print_debug_mode_on(result)
                    pass
                result['Типоразмер'] = '; '.join(indexes)
                result['Название'] = find_system_name()
                print_debug_mode_on(result)
                pass

                all_information = list(dropwhile(lambda x: 'Характеристики входящего оборудования' not in x, self.docx_text))
                print_debug_mode_on(*all_information)

                all_needed_information = []
                all_main_blocks = [piece_of_info for piece_of_info in all_information if piece_of_info.split()[0][-1] == '.']
                main_block_name = [' '.join(piece_of_info.split()[1:]) for piece_of_info in all_main_blocks]
                print_debug_mode_on(all_main_blocks, main_block_name)
                pass
                for i in range(len(all_main_blocks)):
                    clean_information['Название блока'] = all_main_blocks[i]
                    
                    if all_main_blocks[i] != all_main_blocks[-1]:
                        finish_criterion = all_main_blocks[i + 1]
                    else:
                        for possible_finish_criterion in ('Спектральные (дБ) и суммарные (дБА) уровни звуковой мощности', 'Корректированный уровень звукового давления LpA, дБ(А)', 'Примечание:', 'Габаритная схема', 'Габаритные размеры'):
                            if possible_finish_criterion in all_information:
                                finish_criterion = possible_finish_criterion
                                break
                        else:
                            print(all_information)
                            raise ValueError("Очередной странный бланк канального оборудования!")
                    print_debug_mode_on(all_information, all_main_blocks[i], finish_criterion)
                    clean_information['Информация о блоке'] = all_information[all_information.index(all_main_blocks[i])+1 : all_information.index(finish_criterion)]
                    print_debug_mode_on(clean_information['Информация о блоке'])
                    pass
                    
                    clean_information['Информация о блоке'] = '; '.join(clean_information['Информация о блоке'])
                    all_needed_information.append(clean_information)
                    print_debug_mode_on(clean_information)
                    itog_information = result | clean_information
                    print_debug_mode_on(itog_information)
                    ALL_RESULTS.append(itog_information)
                if 'Дополнительное оборудование:' in all_information and 'Габаритная схема' in all_information:
                    clean_information['Название блока'] = 'Дополнительное оборудование'
                    clean_information['Информация о блоке'] = '; '.join(all_information[all_information.index('Дополнительное оборудование:') + 1:all_information.index('Габаритная схема')])
                    itog_information = result | clean_information
                    ALL_RESULTS.append(itog_information)
                print_debug_mode_on(*ALL_RESULTS, sep='\n')
                pass

            elif self.blank_type[self.ready_blank_types[6]]:
                firts_table = tuple(takewhile(lambda x: 'Чиллер серии' not in x, self.docx_text))
                # 'Бланк-заказ', 'Дата бланк-заказа', 'Входящий номер', 'Дата входящего номера', 'Объект', 'Номер объекта', 'Дата', 'Организация', 'Менеджер', 'Выполнил', 'Поток', 'Название', 'Типоразмер'
                print_debug_mode_on(firts_table)
                result['Бланк-заказ'] = firts_table[1].split(' от ')
                result['Дата бланк-заказа'] = result['Бланк-заказ'][1]
                result['Бланк-заказ'] = result['Бланк-заказ'][0]
                result['Входящий номер'] = firts_table[3]
                result['Объект'] = firts_table[7]
                result['Дата'] = firts_table[13]
                result['Организация'] = firts_table[5]
                result['Менеджер'] = firts_table[11]
                result['Выполнил'] = firts_table[9]

                second_table = to_create_table('Предложение / Technical Offer', 'Дополнительные комплектующие (поставляются отдельно)')
                third_table = to_create_table('Состав агрегата / Unit composition', 'Описание агрегата / Unit description')
                fourth_table = to_create_table('Технические характеристики оборудования / Technical data', '(1)	Без учета опций. / Without options')
                print_debug_mode_on(second_table)
                result['Типоразмер'] = second_table[7].split(',')[0]
                print_debug_mode_on(result)
                print_debug_mode_on(third_table, fourth_table, sep='\n')

                # 'Назначение', 'Название блока', 'Тип блока', 'Информация о блоке'
                clean_information['Название блока'] = 'Технические характеристики оборудования'
                clean_information['Информация о блоке'] = fourth_table

                itog_information = result | clean_information
                ALL_RESULTS.append(itog_information)
                pass

            elif self.blank_type[self.ready_blank_types[7]]:
                the_info = to_create_table('Количество отводов', 'Примечание')
                print_debug_mode_on(the_info)

                result['Типоразмер'] = to_create_table('Модель:', 'Примечание')[0].split(' ')[-1].strip()
                print_debug_mode_on(result['Типоразмер'])
                pass

                clean_information['Название блока'] = 'Технические характеристики оборудования'
                clean_information['Информация о блоке'] = the_info

                itog_information = result | clean_information
                ALL_RESULTS.append(itog_information)
                pass
        else:
            itog_information = result | clean_information
            ALL_RESULTS.append(itog_information)

        if self.filename != self.temporary_filename:
            print_debug_mode_on(self.filename, self.temporary_filename)
            pass
            remove(self.temporary_filename)
        else:
            pass
        return result, ALL_RESULTS

class XmlParser:
    """Класс парсера xml данных из файла. Оригинальный класс был создан Саакяном Алексеем, я же его чуть сократил, приспособил под свои нужды"""

    def express_check(self) -> bool:
        """Общая проверка экспресс
        Args:
            xml_root (xml.Tree): xml_root для проверки
        Returns:
            bool: результат проверки
        """
        if len(self.xml_root[0].findall("*[@proIsCustom='usSpecial']")) == 0:
            return all(
                [
                    self.express_name_check(),
                    self.express_block_check(),
                    self.express_coil_check(),
                    self.express_fanfree_check(),
                    self.flanges_insaide(),
                ]
            )
        return False

    def flanges_insaide(self):
        return len(self.xml_root[0].findall("*[@cfnElement='cadCoilsFlangeKit']")) == 0

    def express_name_check(self, **kwargs) -> bool:
        """Функция проверки имени на экспресс
        Args:
            xml_root (xml.etree.ElementTree.Element): xml.etree.ElementTree.Element
            size_list (list, optional): список размеров
        Returns:
            bool: итог проверки
        """
        type_list = kwargs.get("type_list") or ["300", "500"]
        constrution_perfomance = kwargs.get("constrution_perfomance") or {
            "300": ["00"],
            "500": ["01", "02", "03", "04"],
        }
        size_list = kwargs.get("size_list") or [
            "019",
            "039",
            "034",
            "054",
            "058",
            "078",
            "086",
            "087",
            "097",
            "115",
            "116",
            "117",
            "138",
            "156",
            "173",
            "193",
            "194",
            "215",
            "234",
            "240",
        ]
        setup_list = kwargs.get("setup_list") or ["00", "10", "21", "61", "71"]

        features = findall(r"-(\d{3})-(\d{3})-(\d{2})-(\d{2})-([a-zа-яА-ЯЁё0-9]{2,3})-?([В])?", self.name)
        try:
            features = dict(zip(["type", "size", "constrution_perfomance", "setup", "climat", "ex"], features[0]))
            if all(
                [
                    features["type"] in type_list,
                    features["constrution_perfomance"] in constrution_perfomance[features["type"]],
                    features["size"] in size_list,
                    features["setup"] in setup_list,
                    features["ex"] == "",
                    not any([len(features) == 0, "/" in self.name]),
                ]
            ):  # Проверяем все услвия по наименованию
                return True
        except (IndexError, KeyError):
            pass

        return False

    def express_fanfree_check(self, **kwargs) -> bool:
        """Метод проверки ВОСК(ВСК) на экспресс
        Args:
            xml_root (xml.etree.ElementTree.Element): xml.etree.ElementTree.Element
            driver_sizes (list, optional): список допустимых двигателей
            mask (regexp, optional): маска для ВОСК/ВСК
        Returns:
            bool: реззультат проверки
        """
        mask = kwargs.get("mask") or r"ВОСК(\d{2}[Б]?)-(\d{3})-(\d{5}-\d{2})-(\d)-([ОГМ]{1,4})-([А-Я][0-9])"
        type_list = kwargs.get("type_list") or ["62", "72Б", "92"]
        fan_sizes = kwargs.get("fan_sizes") or [
            "023",
            "025",
            "028",
            "032",
            "035",
            "040",
            "045",
            "050",
            "056",
            "063",
            "071",
            "080",
            "090",
        ]
        exc_list = kwargs.get("exc_list") or ["О", "Г", "М"]
        driver_sizes = kwargs.get("driver_sizes") or [
            "00037-02",
            "00055-02",
            "00075-02",
            "00110-02",
            "00150-02",
            "00220-02",
            "00300-02",
            "00400-02",
            "00550-02",
            "00750-02",
            "01100-02",
            "00025-04",
            "00037-04",
            "00055-04",
            "00075-04",
            "00110-04",
            "00150-04",
            "00220-04",
            "00300-04",
            "00400-04",
            "00550-04",
            "00750-04",
            "01100-04",
            "01500-04",
            "01850-04",
            "03000-06",
            "00400-06",
            "00550-06",
            "00750-06",
            "01100-06",
            "01500-06",
        ]
        feat_names = ["type", "size", "driver", "frequency_controller", "exc", "climate_exc"]

        # Получим все вентиляторы ВСК
        fans = set(
            [
                self.xml_root[0].find(k.tag).attrib["proMarking"]
                if self.xml_root[0].find(k.tag).attrib["proMarking"] != ""
                else None
                for k in self.xml_root[0].findall("*[@cfnElement='cadata_frameanFreeAssem']")
            ]
        )
        # Инициализация проверки по движкам (все остальное проверяется в максе
        # регулярке)
        if len(fans) == 0:
            return True

        else:
            for fan in fans:  # Бежим по всем вентиляторам
                reg = dict(zip(feat_names, findall(mask, fan)[0]))
                if all(
                    [
                        reg["type"] in type_list,
                        reg["size"] in fan_sizes,
                        reg["exc"] in exc_list,
                        reg["driver"] in driver_sizes,
                    ]
                ):  # Проверяем условия
                    return True
            return False

    def express_block_check(self, not_express_block_types: list = None) -> bool:
        """Метод проверки блоков на экспресс
        Args:
            xml_root (xml.etree.ElementTree.Element): xml.etree.ElementTree.Element
            not_express_block_types (set, optional): множество блоков которые не проходят проверку
        Returns:
            bool: результат проверки
        """
        not_express_block_types = not_express_block_types or [
            "Fan",
            "SteamHeater",
            "ElectroHeater",
            "CRecoverRecup",
            "RecoverPlate",
            "WasherFor",
            "WasherSot",
            "HRecoverRecup",
            "EInletHeater",
            "EInletEHeater",
            "EInletHeaterEInlet",
            "EInletEHeaterEInlet",
            "SteamHumid",
            "HepaFilter",
            "SorbFilter",
            "CompactFilter",
            "RefrigeratingUnit",
            "EInletUP",
            "UltraViolet",
            "EHeaterTriac",
            "FanMotorWheel",
            "RefriCompUnit",
            "FrCondencer",
            "RefriWaterCond",
            "RecoverRotor",
        ]  # Добавлен RecoverRotor по записке от Горбатенко 230406

        # Проверим все условия на express
        not_express_block_types = set(not_express_block_types)
        block_types = set(
            [
                self.xml_root[0].find(k.tag).attrib["proBlockType"]
                if self.xml_root[0].find(k.tag).attrib["proBlockType"] != ""
                else None
                for k in self.xml_root[0].findall("*[@proBlockType]")
            ]
        ) - set([None])
        return block_types & not_express_block_types == set()

    def express_coil_check(self, **kwargs) -> bool:
        """Метод проверки на экспресс теплообменников
        Args:
            xml_root (xml.etree.ElementTree.Element): xml.etree.ElementTree.Element
            mask (regexp, optional): маска для теплообменники
            coils_type (list, optional): типы теплообменников для поиска
        Returns:
            bool: результат проверки
        """
        mask = (
            kwargs.get("mask")
            or r"[В][Н][В][12][24][3]\.\d-\d{3}-\d{3}-(01|02|03|04)|[В][О][В][23][4][3]\.\d-\d{3}-\d{3}-(03|04|06|08)|[В][О][Ф][23][45][3]\.\d-\d{3}-\d{3}-(03|04|06|08)"
        )
        coils_type = kwargs.get("coils_type") or [
            "cadHeater",
            "cadPHeater",
            "cadCooler",
            "cadata_framerCooler",
            "cadata_framerCondencer",
            "cadHCRecover",
            "cadHCRecover",
            "cadata_framerCooler",
        ]

        # Получим список теплообменников
        coils = []
        for c_type in coils_type:
            coils += [
                self.xml_root[0].find(k.tag).attrib["proMarking"]
                if self.xml_root[0].find(k.tag).attrib["proMarking"] != ""
                else None
                for k in self.xml_root[0].findall(f"*[@cfnElement='{c_type}']")
            ]
        # Проверка через маску регулярки
        for coil in coils:
            reg = findall(mask, coil)
            if not reg:
                return False
        return True

    def define_name(self):
        """Метод определения имени."""
        if (
            "ВЕРОСА" in self.xml_root[0].findall("*[@proUnitName]")[0].attrib["proUnitName"]
            or "Air" in self.xml_root[0].findall("*[@proUnitName]")[0].attrib["proUnitName"]
        ):
            self.name = "Кондиционер " + "/".join(
                list(set([k.attrib["proUnitName"] for k in self.xml_root[0].findall("*[@proUnitName]")]))
            )
        else:
            self.name = "Кондиционер " + "/".join(
                list(set([k.attrib["proFrontName"] for k in self.xml_root[0].findall("*[@proFrontName]")]))
            )

        if len(self.xml_root[0].findall('*[@proOemOrder="oemPAT"]')) != 0:
            self.name = self.name.replace("ВЕРОСА", "PATAIR")

        # Переформатируем название веросы 600
        mask = r"(600)-(\d{3})-(\d{2})-\d{2}-([УХЛ123]{1,4})-?(В)?"
        features = findall(mask, self.name)

        execution_add = {
            "00": 0,
            "02": 20,
            "05": 50,
            "06": 60,
            "07": 70,
            "08": 80,
        }
        explosion_add = {"": 0, "В": 2}

        if len(features) != 0:
            features = features[0]
            core = str(600 + execution_add[features[2]] + explosion_add[features[-1]])
            splited_name = split(mask, self.name)
            self.name = splited_name[0] + "-".join([core, splited_name[2], splited_name[4]])

    def prepare_dict(self):
        """Метод подготовки результирующей структуры для ответа.

        Формат dict
        """
        # Переопределим имя приточной установки
        self.define_name()
        # Подготовка верхних уровней структуры словаря
        self.result_dict[0] = {}
        self.result_dict[0]["name"] = self.name
        self.result_dict[0]["amount"] = 1
        self.result_dict[0]["um"] = None
        self.result_dict[0]["serie"] = self.serie
        self.express = self.express_check()
        self.result_dict[0]["express"] = self.express
        self.result_dict[0]["errors"] = None

        # Пишем папку
        folders = {
            "PATAIR": "00000239856",
            "ВЕРОСА-250": "00000262202",
            "ВЕРОСА-30": "00000238211",
            "ВЕРОСА-50": "00000238212",
            "ВЕРОСА-55": "00000296153",
            "ВЕРОСА-56": "00000307558",
            "ВЕРОСА-6": "00000394495",
            "ВЕРОСА-7": "00000238213",
        }
        self.result_dict[0]["folder"] = None
        for key, value in folders.items():
            if key in self.name:
                self.result_dict[0]["folder"] = value
                break

        self.result_dict[0]["nodes"] = {}

        # Внесем монтажные комплекты
        for element in self.spares_kits + self.montage_kits:
            tag = element.tag
            self.result_dict[0]["nodes"][tag] = {}
            self.result_dict[0]["nodes"][tag]["name"] = (
                self.xml_root[0].find(tag).attrib["cfnName"].replace(self.kit_sign, "").strip().capitalize()
            )
            self.result_dict[0]["nodes"][tag]["amount"] = 1
            self.result_dict[0]["nodes"][tag]["um"] = None
            self.result_dict[0]["nodes"][tag]["serie"] = self.serie
            self.result_dict[0]["nodes"][tag]["nodes"] = {}

    def write_my_kit(self, root, element, element_attr_dict, kits=None):
        """Метод записи комплектов блока и внутреннего каркаса.

        Args:
            root (dict): _description_
            element (EtTree.Element): элемент
            element_attr_dict (dict): словарь элемента
            amount (int, optional): количество. Defaults to 1.
            kits (str, optional): комплекты. Defaults to None.

        Returns:
            _type_: _description_
        """
        # Пишем моноблок
        self.write_element(root, element, element_attr_dict)
        in_root = root[element.tag]["nodes"]
        res = {}
        tag_dict = {"Body": "B", "Inner": "IN"}

        name_dict = {"Body": "Корпус ", "Inner": "Внутренний каркас  "}
        for kit in kits.split(" "):
            # Пишем корпус в моноблок
            in_root[element.tag + tag_dict[kit]] = {}
            in_root[element.tag + tag_dict[kit]]["name"] = name_dict[kit] + element_attr_dict["cfnName"]
            in_root[element.tag + tag_dict[kit]]["amount"] = float(element.attrib["cfnAmount"])
            in_root[element.tag + tag_dict[kit]]["um"] = (
                element_attr_dict["cfnNote"].replace("^", "")
                if (
                    element_attr_dict["cfnNote"] != ""
                    and len(element_attr_dict["cfnNote"]) <= 6
                    and element_attr_dict["cfnNote"][0].isalpha()
                )
                else None
            )
            #  and element_attr_dict['cfnNote'][0].isalpha()) else 'шт'
            in_root[element.tag + tag_dict[kit]]["serie"] = self.serie
            in_root[element.tag + tag_dict[kit]]["buy"] = 1 if element_attr_dict["cfnLevel"] in self.buy_list else 0
            in_root[element.tag + tag_dict[kit]]["nodes"] = {}

            res[kit] = in_root[element.tag + tag_dict[kit]]["nodes"]

        return res

    def recur_kit_finder(self, parent):
        """Метод поиска элементов, входящих в комплекты.

        Args:
            parent (EtTree.Element): родитель для поиска комплектов
        """
        for element in parent:
            self.get_kit_for_element(element)
            self.recur_kit_finder(element)

    def recur_xml_to_dict(self, **kwargs):
        """Метод рекурсивного обхода структуры с сбором элементов и
        подстановкой уже полученнхы с других ресурсов структур в результирующем
        словаре."""
        amount = float(kwargs.get("amount") or 1)
        parent = kwargs["parent"]
        dont_sum_elements = kwargs.get("dont_sum_elements", None)

        for element in parent:
            # pdmo(parent, element)
            if self.get_kit_for_element(parent=element):
                # Фиксируем параметры про-родителя и ребенка
                element_attr_dict = self.xml_root[0].find(element.tag).attrib
                parent_attr_dict = self.xml_root[0].find(parent.tag).attrib

                if len(element) == 0 and element_attr_dict["cfnElement"] == "cadJointFolder":
                    return

                # Проверяем содержится ли тип эемента в тех, которые надо выделить
                # корпус и внутренние каркасы
                if any(
                    [
                        element_attr_dict["cfnLevel"] in self.body_types,
                        element_attr_dict["cfnElement"] in self.body_types,
                    ]
                ):
                    root = kwargs["body_root"] if kwargs.get("body_root", None) is not None else kwargs["root_dict"]

                elif (
                    any(
                        [
                            element_attr_dict["cfnLevel"] in self.inner_types,
                            element_attr_dict["cfnElement"] in self.inner_types,
                        ]
                    )
                    and parent_attr_dict["cfnElement"] != "cadJointFolder"
                ):
                    root = kwargs["inner_root"] if kwargs.get("inner_root", None) is not None else kwargs["root_dict"]

                else:
                    root = kwargs["root_dict"]

                # Если попали в игнорируемый список, проваливаемся ниже
                if any(
                    [
                        element_attr_dict["cfnLevel"] in self.ignore_types,
                        element_attr_dict["cfnElement"] in self.ignore_types,
                        element.tag in [elem.tag for elem in self.montage_kits + self.spares_kits],
                    ]
                ):
                    self.recur_xml_to_dict(
                        parent=element,
                        root_dict=root,
                        body_root=kwargs.get("body_root", None),
                        inner_root=kwargs.get("inner_root", None),
                    )

                # Если  разузловываем тип, заберем спецухи и вставим в нашу
                # структуру
                elif element_attr_dict["cfnElement"] in self.unravel_types:
                    self.write_element(root, element, element_attr_dict)
                    if self.structure_targets.get(element.tag):  # Проверим наличие а найденных
                        root[element.tag]["nodes"] = self.structure_targets.get(element.tag)
                        self.recur_kit_finder(element)

                    # Иначе читаем то, что есть в xml игнорируя те, которые
                    # холодильные
                    elif element_attr_dict["cfnElement"] not in [
                        "cadRefriCompUnit",
                        "cadRefriCoilsUnit",
                        "cadRefriCondUnit",
                        "cadRefriWaterCondUnit",
                    ]:
                        self.recur_xml_to_dict(
                            parent=element,
                            root_dict=root[element.tag]["nodes"],
                            body_root=kwargs.get("body_root", None),
                            inner_root=kwargs.get("inner_root", None),
                            dont_sum_elements=True,
                        )

                # Если наткнулись на моноблок, то проверяем надо ли делать
                # дополнительную раскладку
                elif element_attr_dict["cfnElement"] == "cadMonoblockFolder":
                    if len(self.body_types) != 0 and len(self.inner_types) != 0:
                        roots = self.write_my_kit(root, element, element_attr_dict, kits="Body Inner")
                        body_root, inner_root = roots["Body"], roots["Inner"]
                    elif len(self.body_types) == 0 and len(self.inner_types) != 0:
                        roots = self.write_my_kit(root, element, element_attr_dict, kits="Inner")
                        body_root, inner_root = None, roots["Inner"]
                    elif len(self.body_types) != 0 and len(self.inner_types) == 0:
                        roots = self.write_my_kit(root, element, element_attr_dict, kits="Body")
                        body_root, inner_root = roots["Body"], None
                    else:
                        self.write_element(root, element, element_attr_dict)
                        body_root, inner_root = None, None

                    self.recur_xml_to_dict(
                        parent=element,
                        root_dict=root[element.tag]["nodes"],
                        body_root=body_root,
                        inner_root=inner_root,
                        amount=None,
                    )

                # Заберем элементы сруктуры, которые запрашивали с других ресурсов,
                # но не трубется записывать в итоговую спецификацию
                # elif element_attr_dict["cfnElement"] in self.query_prod_unhead + self.unravel_types:
                #     if self.structure_targets.get(element.tag):  # Проверим наличие а найденных
                #         self.add_specification_from_source(root, element.tag)
                #     else:  # Иначе читаем то, что есть в xml
                #         self.recur_xml_to_dict(
                #             parent=element,
                #             root_dict=root,
                #             body_root=kwargs.get("body_root"),
                #             inner_root=kwargs.get("inner_root"),
                #         )

                else:  # Пишем элементы в структуру, если не сработали другие условия
                    self.write_element(root, element, element_attr_dict, dont_sum_elements)
                    self.recur_xml_to_dict(
                        parent=element,
                        root_dict=root[element.tag]["nodes"],
                        body_root=kwargs.get("body_root"),
                        inner_root=kwargs.get("inner_root"),
                        amount=None,
                    )

    def write_element(self, root, element, element_attr_dict, dont_sum_elements=None):
        """Метод записи элемента в результирующий словарь.

        Args:
            root (dict): корень в результирующем словаре для записи
            element (EtTree.Element): элемент, который надо записать
            element_attr_dict (dict): словарь свойств элемента
            amount (int, optional): прокинутый множитель количества. Defaults to 1.
        """

        if element.tag in root:  # Если есть такой же элемент, накинем количество
            if not dont_sum_elements:
                root[element.tag]["amount"] += float(element.attrib["cfnAmount"])
        else:  # Если нет такого же, то делаем новую запись
            root[element.tag] = {}
            root[element.tag]["name"] = (
                element_attr_dict["cfnName"] + " " + element_attr_dict["cfnDesignation"]
                if element_attr_dict["cfnDesignation"] != ""
                else element_attr_dict["cfnName"]
            )

            # ! Убирать костылинг!!!
            # ! Добавлено по просьбе Горбатенко (231027)
            # TODO: Удалить после перехода продацов на 50+ версию ВЕРОСА
            if ' ГОСТ 9941-81' in  root[element.tag]["name"]:
                root[element.tag]["name"] = root[element.tag]["name"].replace(
                    ' ГОСТ 9941-81', ' ГОСТ 9941-2022')

            # ! ДОБАВЛЕНО ПО ПРОСЬБЕ ТРУСОВ - убрать потом (230505)
            if all([element_attr_dict["cfnElement"] == "cadUntyped", "-5.6-" in root[element.tag]["name"]]):
                root[element.tag]["name"] = root[element.tag]["name"].replace("-5.6-", "-8.8-")

            # ! ДОБАВЛЕНО ПО ПРОСЬБЕ ТРУСОВ - убрать потом (230705)
            if all([element_attr_dict["cfnElement"] == "cadUntyped", "-5-А9В" in root[element.tag]["name"]]):
                root[element.tag]["name"] = root[element.tag]["name"].replace("-5-А9В", "-8-А9В")

            # ! ДОБАВЛЕНО ПО ПРОСЬБЕ ТРУСОВ - убрать потом (231003)
            if element_attr_dict["cfnElement"] == "cadEMotor":
                root[element.tag]["name"] = (
                    root[element.tag]["name"]
                    .replace("IM1001", "IM1081")
                    .replace("IM1031", "IM1081")
                    .replace("IM1011", "IM1081")
                )
                root[element.tag]["name"] = root[element.tag]["name"].replace("220/380", "380")
                root[element.tag]["name"] = root[element.tag]["name"].replace("-F-", "-0-")
                root[element.tag]["name"] = root[element.tag]["name"].replace("-У3-", "-У2-")

            # ! ДОБАВЛЕНО ПО ПРОСЬБЕ ТРУСОВ - убрать потом (231103)
            if element_attr_dict["cfnElement"] == "cadMaterial" and\
                'Прокат' in root[element.tag]["name"]:
                root[element.tag]["name"] = root[element.tag]["name"].replace('-Ц140','-Ц275')
            # ! Конец

            print_debug_mode_on(element_attr_dict)

            root[element.tag]["amount"] = float(element.attrib["cfnAmount"])
            root[element.tag]["um"] = (
                element_attr_dict["proMeasureUnitStr"].replace("^", "")
                    if "proMeasureUnitStr" in element_attr_dict else None
            )
            # and element_attr_dict['cfnNote'][0].isalpha()) else 'шт'
            root[element.tag]["serie"] = None
            root[element.tag]["buy"] = 1 if element_attr_dict["cfnElement"] in self.buy_list else 0
            root[element.tag]["nodes"] = {}
            for key, value in element_attr_dict.items():
                root[element.tag][key] = value

    def add_specification_from_source(self, root, tag):
        """Метод добавления найденных составов для элементов tag, в указанное
        место результирующего словаря root с множителем amount.

        Args:
            root (dict): место в словаре (корень)
            tag (str): ИД элемента
            amount (int, optional): прокинутое количество. Defaults to 1.
        """
        for key in self.structure_targets[tag]:
            name = self.structure_targets[tag][key]["name"] if self.structure_targets[tag][key]["name"] else ""
            new_key = "N" + str(uuid3(NAMESPACE_DNS, name)).replace("-", "")
            if new_key in root:
                try:
                    assert root[new_key]["name"] == deepcopy(self.structure_targets[tag][key]["name"])
                    root[new_key]["amount"] += deepcopy(self.structure_targets[tag][key]["amount"])
                except AssertionError:
                    print("Assertion arror!!!")
                    print(self.structure_targets[tag])
                    print(self.structure_targets[tag][key]["name"])
            else:
                root[new_key] = deepcopy(self.structure_targets[tag][key])

    def get_kit_for_element(self, parent):
        """Метод сборка элементов из комплектов. Используется как идентификатор
        для наличия действия помещения в комплект в родителе. Возвращает True
        если не входит в комплект, False если входит.

        Args:
            parent (EtTree.Element): элемент для проверки на запись в комплекты
        """
        for element in parent:
            # pdmo(parent, element)
            if element.tag in [elem.tag for elem in self.spares_kits + self.montage_kits]:
                kit_root = self.result_dict[0]["nodes"][element.tag]["nodes"]
                par_attr_dict = self.xml_root[0].find(parent.tag).attrib
                self.write_element(kit_root, parent, par_attr_dict)
                return False
        return True

    def clear_result(self, input_dict: dict, level=2):
        """Метод очистки результирующего словаря от пустых элементов.

        Args:
            input_dict (dict): входной словарь
            level (int, optional): глубина очистки. Defaults to 1.
        """
        if level == 0:
            return
        else:
            keys = list(input_dict["nodes"].keys())
            for key in keys:
                if all(
                    [
                        len(input_dict["nodes"][key]["nodes"]) == 0,
                        any(
                            [
                                "крепления установки" in input_dict["nodes"][key]["name"],
                                "Корпус Моноблок" in input_dict["nodes"][key]["name"],
                                "Внутренний каркас" in input_dict["nodes"][key]["name"],
                                "Монтажный комплект" in input_dict["nodes"][key]["name"],
                            ]
                        ),
                    ]
                ):
                    input_dict["nodes"].pop(key)
                else:
                    self.clear_result(input_dict["nodes"][key], level - 1)

    def def_crates(self, monoblocks: list) -> dict:
        """Определение обрешетки

        Args:
            monoblocks (list): список моноблоков

        Returns:
            dict: структура упаковки
        """

        res = {} # Начальное состояние

        for monoblock in monoblocks:
            length = int(monoblock.attrib['proLength']) + 50
            width = int(monoblock.attrib['proWidth']) + 100
            height = int(monoblock.attrib['proHeight']) + 50

            amt_count = ((ceil(width/150 if width<=1000 else width/160)*length)\
                    + (height + 125)*(0 if (height + 125)==0 else (8+0 if length<=800 \
                        else (2 if length<=2000 else (4 if length<=2500 else (6 if \
                            length<=3000 else 0)))) + (0 if width<=800 else (2 if \
                                width<=2000 else (4 if width<=2500 else (6 if \
                                    width<=3000 else 0)))))\
                    + (height + 50)*2 + (width + 100) * 2\
                    + (0 if height == 0 else (length + 50) if height < 1000 else \
                        ((length ** 2) + (height + 25) ** 2) ** 0.5) * 2\
                    + (0 if height==0 else (width + 100) if height<1000 else \
                        sqrt((width-25)**2+(height+25)**2)) * 2\
                    + (width + 100) * (2 if length <= 1000 else 3 if length <= 2000 \
                    else 4 if length <= 2500 else 5)) / 1000 * 0.1 * 0.025

            if 'P00001' in res:
                res['P00001']['amount'] += round_up(amt_count, 3)
            else:
                res['P00001'] = {
                    'name': 'Доска-2 хв-25х100х6000 ГОСТ 8486-86',
                    'amount': round_up(amt_count, 3),
                    'um': None,
                    'serie': None,
                    'buy': 0,
                    'nodes': {}
                }


            amt_count = width * (2 if length <= 1000 else 3 if length <= 2000 \
                else 4 if length <= 2500 else 5) / 1000 * 0.1 * 0.05

            if 'P00002' in res:
                res['P00002']['amount'] += round_up(amt_count, 3)
            else:
                res['P00002'] = {
                    'name': 'Брус-2 хв-50х100х6000 ГОСТ 8486-86',
                    'amount': round_up(amt_count, 3),
                    'um': None,
                    'serie': None,
                    'buy': 0,
                    'nodes': {}
                }

        return res

    def def_plywood_box(self, monoblocks: list) -> dict:
        """Определение фанерного ящика

        Args:
            monoblocks (list): список моноблоков

        Returns:
            dict: структура упаковки
        """

        res = {} # Начальное состояние

        for monoblock in monoblocks:
            length = int(monoblock.attrib['proLength']) + 50
            width =  ((int(monoblock.attrib['proWidth'])+100)//100+1) if \
                ((int(monoblock.attrib['proWidth'])+100)%100)>0 else \
                    ((int(monoblock.attrib['proWidth'])+100)//100)*100
            height = int(monoblock.attrib['proHeight']) + 50

            amt_count = ((width* width/100) + ((width + 50) * 2) + ((width - 150) * 2)\
                + (height + 125) * (8 + 0 if length <= 1000 else \
                    2 if length <= 2000 else 4 if length <= 2500 else \
                        6 + 0 if width <= 1000 else 2 if width <= 2000 \
                            else 4 if width <= 2500 else 6)\
                + 2 * (width + 50) *  (2 if length <= 1000 else 3 if length <= 2000 \
                    else 4 if length <= 2500 else 5)) / 1000 * 0.1 * 0.025

            if 'P00001' in res:
                res['P00001']['amount'] += round_up(amt_count, 3)
            else:
                res['P00001'] = {
                    'name': 'Доска-2 хв-25х100х6000 ГОСТ 8486-86',
                    'amount':  round_up(amt_count, 3),
                    'um': None,
                    'serie': None,
                    'buy': 0,
                    'nodes': {}
                }

            amt_count = (2 if length <= 1000 else 3 if length <= 2000 else 4 if length <= 2500 else 5) * width / 1000 * 0.1 * 0.05

            if 'P00002' in res:
                res['P00002']['amount'] += round_up(amt_count, 3)
            else:
                res['P00002'] = {
                    'name': 'Брус-2 хв-50х100х6000 ГОСТ 8486-86',
                    'amount': round_up(amt_count, 3),
                    'um': None,
                    'serie': None,
                    'buy': 0,
                    'nodes': {}
                }

            if 'P00003' in res:
                res['P00003']['amount'] += 5

            else:
                res['P00003'] = {
                    'name': 'Фанера береза/береза, ФСФ, III/III, Е1, НШ, 2440х1220х6.5, 5 слоёв ГОСТ 3916.1-96',
                    'amount': 5,
                    'um': None,
                    'serie': None,
                    'buy': 0,
                    'nodes': {}
                }

        return res

    def def_wooden_boxes(self, monoblocks: list) -> dict:
        """Определение деревянного ящика

        Args:
            monoblocks (list): список моноблоков

        Returns:
            dict: структура упаковки
        """

        res = {} # Начальное состояние

        for monoblock in monoblocks:
            length = (ceil(int(monoblock.attrib['proLength'])/100) + 1 ) * 100
            width = (ceil(int(monoblock.attrib['proWidth'])/100) + 1 ) * 100
            height = int(monoblock.attrib['proHeight']) + 50

            amt_count = (length * width / 100 + (height + 125) * (length / 100 * 2)\
                + (height + 125) * (width / 100 * 2) + (width + 50) * \
                    (length / 100 * 2)) / 1000 * 0.1 * 0.025

            if 'P00001' in res:
                res['P00001']['amount'] += round_up(amt_count, 3)
            else:
                res['P00001'] = {
                    'name': 'Доска-2 хв-25х100х6000 ГОСТ 8486-86',
                    'amount': round_up(amt_count, 3),
                    'um': None,
                    'serie': None,
                    'buy': 0,
                    'nodes': {}
                }

            amt_count = width * (2 if length <= 1000 else 3 if length <= 2000 else 4 \
                if length <= 2500 else 5) / 1000 * 0.1 * 0.05

            if 'P00002' in res:
                res['P00002']['amount'] += round_up(amt_count, 3)
            else:
                res['P00002'] = {
                    'name': 'Брус-2 хв-50х100х6000 ГОСТ 8486-86',
                    'amount': round_up(amt_count, 3),
                    'um': None,
                    'serie': None,
                    'buy': 0,
                    'nodes': {}
                }

        return res

    def def_sea_box(self, monoblocks: list) -> dict:
        """Определение морского ящика

        Args:
            monoblocks (list): список моноблоков

        Returns:
            dict: структура упаковки
        """

        res = {} # Начальное состояние

        for monoblock in monoblocks:

            length = (ceil(int(monoblock.attrib['proLength']) / 100) + 1 ) * 100
            width = (ceil(int(monoblock.attrib['proWidth']) / 100) + 1 ) * 100
            height = int(monoblock.attrib['proHeight']) + 50

            amt_count = ((ceil(length / 100) * 2 + 2) * (height + 125) \
                + (ceil(length / 100) * 2) * (height + 125) + (width + 50) * \
                    ceil(length / 100)) / 1000 * 0.1 * 0.025

            if 'P00001' in res:
                res['P00001']['amount'] += round_up(amt_count, 3)
            else:
                res['P00001'] = {
                    'name': 'Доска-2 хв-25х100х6000 ГОСТ 8486-86',
                    'amount': round_up(amt_count, 3),
                    'um': None,
                    'serie': None,
                    'buy': 0,
                    'nodes': {}
                }

            amt_count = ((length / 12 * width) + 6 * 200 + 2 * (length + 150) + 4 * length \
                + 4 * (width - 100) + 4 * (height - 100) + width * 2) / 1000 * 0.1 * 0.05

            if 'P00002' in res:
                res['P00002']['amount'] += round_up(amt_count, 3)
            else:
                res['P00002'] = {
                    'name': 'Брус-2 хв-50х100х6000 ГОСТ 8486-86',
                    'amount': round_up(amt_count, 3),
                    'um': None,
                    'serie': None,
                    'buy': 0,
                    'nodes': {}
                }

        return res

    def define_packing(self, packing, for_sales=False):
        """Определяем упаковку для моноблоков

        Args:
            packing (str): упаковка
        """
        monoblocks = self.xml_root[0].findall("*[@cfnElement='cadMonoblockFolder']")

        if packing == 'Обрешетка':
            res = self.def_crates(monoblocks)

        elif packing == 'Ящик из фанеры':
            res = self.def_plywood_box(monoblocks)

        elif packing == 'Сплошной ящик':
            res = self.def_wooden_boxes(monoblocks)

        elif packing == 'Морская упаковка':
            res = self.def_sea_box(monoblocks)

        else:
            res = {}
            msg = '\nВыбранный тип упаковки не поддерживается для этого изделия.'
            self.result_dict[0]['errors'] = self.result_dict[0]['errors'] + msg \
                    if self.result_dict[0]['errors'] else msg

        if for_sales: # Проверяем ключ для продаж или нет
            self.result_dict[0]['nodes']['PACKING'] = {
                'name': 'Упаковка',
                'amount': 1,
                'um': None,
                'serie': self.serie,
                'nodes': res
            }

        else:
            for key, item in res.items():
                self.result_dict[0]['nodes'][key] = item

    def __init__(self, file_path, **kwargs):
        """Метод инициализации.

        Args:
            file_path (str): путь к файлу для работы парсера
        """
        # Прокидываем агрументы из kwargs и пишем переменные класса
        spare_kit_feat = kwargs.get("spare_kit_feat") or "cadQualContrAddToKitSpares"
        montage_kit_feat = kwargs.get("montage_kit_feat") or "cadQualContrAddToKitMontage"
        self.xml_root = ET.parse(file_path).getroot()  # Грузим файл

        # Комплекты
        self.kits = kwargs.get("kits") or ["cadQualContrAddToKitMontage", "cadQualContrAddToKitSpares"]
        # Ищем все комплекты
        self.spares_kits = self.xml_root[0].findall(f"*[@cfnElement='{spare_kit_feat}']")
        self.montage_kits = (
            self.xml_root[0].findall(f"*[@cfnElement='{montage_kit_feat}']")
            if "Внутренние карскасы" not in self.kits
            else []
        )

        # Подготоваливаем переменные для записи
        self.kit_sign = "Поместить в "
        self.result_dict = {}
        self.targets = []
        self.structure_targets = {}
        self.n_bump_stop = 0
        self.errors = ""
        self.name = None
        # Читаем настройки запроса
        self.plan_spec = kwargs.get("plan_spec") or False
        self.buy_list = kwargs.get("buy_list") or []
        # Типы, в которые не пристыковывать обозначение
        self.designation_unstack = kwargs.get("designation_unstack") or []

        # Типы для игнорирования
        self.ignore_types = []
        # self.ignore_types = kwargs.get("ignore_types") or [
        #     # Уровени cfnLevel из XML Рябова
        #     "caeCubGeom",
        #     "caePropCollection",
        #     "caePropCollectionItem",
        #     "cadMonoblocksFolder",
        #     "cadUnitFolder",
        #     "cadUnitsFolder",
        #     "cadBlocksFolder",
        #     "cadInnercarcassFolder",
        #     "cadBlockFolder",
        #     "cadPanel",
        #     "cadOutercarcassFolder",
        #     "cadTimeLine",
        #     "cadTriangl",
        #     "cadSupport",
        #     "cadSupportDetail",
        #     "cadSupportFolder",
        #     "cadQualityControlItem",
        #     "cadCoilsFin",
        #     "cadConnAct",
        #     "cadata_framean",
        #     "cadJointFolder",
        #     "cadHeaterDD",
        #     "cadFrCoolerDD",
        #     "cadCoolerDD",
        #     # Последние 4 элемента удалить после тестов
        #     "cadata_framerCoolerDD",
        #     "cadRigel",
        #     "cadPanelAssem",
        #     "cadRact",
        #     "cadSupportAssem",
        # ]

        # Типы для разузловки
        self.unravel_types = []
        # self.unravel_types = kwargs.get("unravel_types") or [
        #     "cadHeater",
        #     "cadata_frameanAssem",
        #     "cadEHeater",
        #     "cadPHeater",
        #     "cadFrCooler",
        #     "cadata_framerCooler",
        #     "cadFanFreeAssem",
        #     "cadFrCooler",
        #     "cadNoice",
        #     "cadLightSwitchKit",
        #     "cadCoilsFlangeKit",
        #     "cadConnAssem",
        #     "Выкатная деталь ДС",
        #     "Плита подмоторная Брянск",
        #     "Коллектор КАНБАН",
        #     "Шильд",
        #     "Исп механ КПД-4-01",
        #     "Основа",
        #     "Панель шумоглушителя СТАМ",
        #     "Трубка распределителя",
        #     "Патрубок Холод",
        #     "Заготовка распределителя",
        #     "Вал",
        #     "Уголок прокат",
        #     "Выкатная деталь Ф",
        #     "Стенка шумоглушителя",
        #     "Ниппель",
        #     "Трубка",
        #     "Отвод",
        #     "Калач",
        #     "Патрубок",
        #     "Коллектор",
        #     "Капилляр",
        #     "Диск",
        #     "Выкатная деталь",
        #     "Дюза",
        #     "Поддон кооперация",
        #     "Выкатная деталь ДС",
        #     "Втулка для СБ Рк",
        #     "Выкатная деталь ДИП",
        #     "Колесо рабочее сварное",
        #     "Теплообменник",
        #     "Корпус МАКК",
        #     "Вентилятор для ТО",
        #     "Шкаф управления",
        #     "Корпус",
        #     "Колесо рабочее сборное",
        #     "Стакан монтажный",
        #     "Крыльчатка",
        #     "Натяжное устройство",
        # ]

        if "Корпуса" in self.kits:  # Проверяем запросы на раскладку корпусов и внутренних каркасов
            self.body_types = [
                "cadOutercarcassFolder",
                "cadSupportFolder",
                "cadCarcassProtectivBush",
            ]
        else:
            self.body_types = []

        if "Внутренние каркасы" in self.kits:
            self.inner_types = ["cadInnercarcassFolder", "cadGroundBus"]  # 'caeStandard', 'cadUntyped',
            self.ignore_types += ["cadQualContrAddToKitMontage"]
        else:
            self.inner_types = []

        # packing = kwargs.get('packing', None)

        # Типы узлов, которые надо запрашивать в ПДМ, но не выводить их
        # головное имя
        # self.query_prod_unhead = kwargs.get("query_prod_unhead") or ["cadInnercarcassAssem", "cadRoofOverhang"]
        # self.series_types = kwargs.get("series_types") or ["cadMonoblockFolder"]  # Типы для пристыковки серий
        # Находим номер бланка
        self.serie = self.xml_root[0].find(self.xml_root[1][0][0].tag).attrib["cfnName"].strip().replace("Заказ ", "")

        # Подготовим конечную структуру для записи. Запишем голову и другие
        # параметры
        self.prepare_dict()

        if "Air" in self.name:  # Заглушка для загрузки аэрмейта целиком
            self.body_types = []
            self.inner_types = []

        self.recur_xml_to_dict(
            parent=self.xml_root[1][0][0], root_dict=self.result_dict[0]["nodes"]
        )  # Рекурсивно проходим xml и записываем в json

        # self.result_dict[0]["errors"] = self.errors if not kwargs["plan_spec"] else None

        # Читстим от пустых элементов верхнего уровня результирующий словарь
        if "Air" not in self.name:
            self.clear_result(self.result_dict[0])

        # if packing:
        #     self.define_packing(packing, "Корпуса" in self.kits)

        self.result_dict[0]["production_cycle"] = 9  # Определеяем производственный цикл

        # self.json = dumps(self.result_dict[0])  # Сбрасываем в json

class XmlParserForAutomata(XmlParser):
    """Сверху над классом, который создал Саакян, я ещё наворачиваю класс, который ещё добавляет один самый важный для меня словарь-представление данных и одну самую важную (по моим расчётам) функцию - поиск в нужном блоке нужного параметра
    """

    def __init__(self, file_path, **kwargs):
        """Та же самая инициализация, только добавляется атрибут short_data - просто блоки без лишней мишуры. И ещё много что другое

        Args:
            file_path (_type_): адрес файла
        """

        def recur_reduction(data:dict):
            """Рекурсивное сокращение гигантского словаря: извлекает в short_data только то, что мне требуется для определения комплекта автоматики. Просто рекурсивно обходит словарь и, находя нужное, сразу записывает, ничего не возвращая

            Args:
                data (dict): текущий словарь
            """
            for key in data.keys():
                # print(key, value)
                if any(the_word in data[key]['name'] for the_word in ('Блок ', 'Коллекция элементов отчета проект', 'Коллекция элементов отчета установка')):
                    self.short_data[data[key]['name']] = data[key]['nodes']
                if data[key]['nodes'] != {}:
                    recur_reduction(data[key]['nodes'])

        super().__init__(file_path, **kwargs)
        self.short_data = {}
        recur_reduction(self.result_dict)

        self.main_information = {key: self.find_value_by_name('Коллекция элементов отчета проект', value) for key, value in zip(Blank.all_columns[0:10], (  # 'Бланк-заказ', 'Дата бланк-заказа', 'Входящий номер', 'Дата входящего номера', 'Объект', 'Номер объекта', 'Дата', 'Организация', 'Менеджер', 'Выполнил'
            'Элемент отчета ptgORDER.ptPRJREF',  # Бланк-заказ
            'Элемент отчета ptgIN.ptInDateTime',  # Дата бланк-заказа
            'Элемент отчета ptgIN.ptInRef',  # Входящий номер
            'Элемент отчета ptgIN.ptInDateTime',  # Дата входящего номера
            'Элемент отчета ptgORDER.ptPRJWHERE',  # Объект
            'Элемент отчета ptgIN.ptInRef',  # Номер объекта
            'Элемент отчета ptgORDER.ptDATETIME',  # Дата
            'Элемент отчета ptgCUSTOMER.ptCOMPANY',  # Организация
            'Элемент отчета ptgADMINISTRANT.ptUSER0',  # Менеджер
            'Элемент отчета ptgADMINISTRANT.ptUSER',  # Выполнил
        ))} | {key: self.find_value_by_name('Коллекция элементов отчета установка', value) for key, value in zip(Blank.all_columns[10:-4], (  # 'Поток', 'Название', 'Типоразмер'
            'Элемент отчета ptgPARAMETERS.ptAIRDirect',  # Поток
            'Элемент отчета ptgPARAMETERS.ptUnitName',  # Название
            'Элемент отчета ptgPARAMETERS.ptSize'  # Типоразмер
        ))}  # Колонки основной информации

        self.main_information = {key: value['proReportPropValue'] if value is not None else '-' for key, value in self.main_information.items()}

        self.ALL_MAIN_INFO = {key: self.short_data[key] for key in tuple(self.short_data.keys())[2:]}

        self.IS_OTHERS, self.IS_VEROSA, self.IS_OBPROM, self.IS_CHANAL, self.IS_INDUST, self.IS_INTEPU = False, True, False, False, False, False

    def find_value_by_name(self, adress:str, name:str, multiple_result=False, is_in=False) -> dict | list:
        """Нахождение значения по имени и адресу

        Args:
            adress (str): блок, в котором мы ищем нужную информацию
            name (str): название нужной информации согласно руководству Рябова
            multiple_result (bool, optional): _description_. Defaults to False.
            is_in (bool, optional): _description_. Defaults to False.

        Returns:
            dict: словарь с нужной информацией
        """

        def recur_find(data:dict) -> dict|None:
            """Рекурсивный обход словаря

            Args:
                data (dict): текущий словарь

            Returns:
                dict|None: найденное значение или ничего - если ничего, то ищем дальше и не возвращаем
            """
            for key, value in data.items():
                if data[key]['nodes'] != {}:
                    result = recur_find(data[key]['nodes'])
                    if result is not None:
                        return result
                if criterium(value):
                    return value
        
        criterium = lambda x: (name in x['name'] if is_in else x['name'] == name)

        all_result = []
        if multiple_result:
            for the_value in self.short_data[adress].values():                
                # brbrdeng = list()
                # print(brbrdeng)
                all_result += [val for val in the_value['nodes'].values() if criterium(val)]
                # print(all_result)
                pass
            return all_result
        else:
            return recur_find(self.short_data[adress])

    def print_data(self, full_data_name='full_data.json', short_data_name='short_data.json', print_full_data=False, print_short_data=True):
        """Печать в json словарей для проверки

        Args:
            full_data_name (str, optional): Адрес файла с полной информацией. По умолчанию записывается в 'full_data.json', лежащий в той же папке, что и программа.
            short_data_name (str, optional): Адрес файла с сокращённой информацией. По умолчанию записывается в 'short_data.json', лежащий в той же папке, что и программа.
            print_full_data (bool, optional): Надо ли вообще выводить полную информацию (можно выводить либо что-то одно, либо оба сразу - если оба будут ложью, то нафига вообще эту функцию вызывать?). По умолчанию False (это гигантский словарь в 20 тысяч строк, занимающий около 2 мегабайт, поэтому думайте сами).
            print_short_data (bool, optional): Надо ли вообще выводить сокращённую информацию (можно выводить либо что-то одно, либо оба сразу - если оба будут ложью, то нафига вообще эту функцию вызывать?). По умолчанию True.
        """

        with open(full_data_name, 'w', encoding='utf-8') as file_0, open(short_data_name, 'w', encoding='utf-8') as file_1:
            if print_full_data:
                dump(self.result_dict, file_0, ensure_ascii=False, indent=4)
            if print_short_data:
                dump(self.short_data, file_1, ensure_ascii=False, indent=4)

class DocxExpand:
    """Дополнительные методы или функции, которые должны расширить функционал модуля docx, в основном тут заимствованный у чужих код
    """

    @staticmethod
    def delete_paragraph(paragraph):
        """Удаляет выбранный параграф

        Args:
            paragraph (_type_): объект-параграф
        """
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    
    @staticmethod
    def add_paragraph_after_paragraph(previous_paragraph, new_paragraph_text:str, document):
        """_summary_

        Args:
            previous_paragraph (_type_): объект-параграф
            new_paragraph_text (str): текст нового параграфа
            document (_type_): объект-документ

        Returns:
            _type_: new_paragraph, нужно, чтобы иметь возможность добавить что-то после него.
        """
        new_paragraph = document.add_paragraph(new_paragraph_text)
        previous_paragraph._p.addnext(new_paragraph._p)
        return new_paragraph

    @staticmethod
    def add_paragraph_before_table(paragraph_text:str, table, document):
        """Добавляет параграф перед таблицей

        Args:
            paragraph_text (str): текст нового параграфа
            table (_type_): объект-таблица
            document (_type_): объект-документ
        """
        table._element.addprevious(document.add_paragraph(paragraph_text)._p)

    @staticmethod
    def add_paragraph_after_table(table):
        """Добавляет параграф после таблицы

        Args:
            table (_type_): объект-таблица

        Returns:
            Paragraph: объект-параграф (так надо)
        """
        return Paragraph(table._tbl.getnext(), table._parent).insert_paragraph_before()
    
    @staticmethod
    def move_table_after_paragraph(table, paragraph):
        """Передвигает таблицу после параграфа

        Args:
            table (_type_): объект-таблица
            paragraph (_type_): объект-параграф
        """
        paragraph._p.addnext(table._tbl)
    
    @staticmethod
    def remove_row(table, row):
        """Удаляет выбранную строку данной таблицы

        Args:
            table (_type_): объект-таблица
            row (_type_): объект-строка данной таблицы
        """
        table._tbl.remove(row._tr)
    
    @staticmethod
    def change_cell_applying_its_style(cell, text:str):
        """Недавно обнаруженный костыль выявил занятный факт - внутри ячеек таблицы находятся параграфы, и их можно извлечь как список параграфов, и если не передавать тексту параграфа новое значение, а просто добавить новое, то тогда сохранится исходный стиль, который был у текста в ячейке. Важно - сохраняется стиль, а если к самому стилю были применены изменения внутри шаблона, то эти изменения не сохранятся. То есть, если внутри ячейки текст стиля "Нормальный", стиль "Нормальный" - 12 кегль, во всю ширину, но в шаблоне в ячейке он сделан жирным, жирнота спадёт. Отсюда совет - можно создавать бесконечно много стилей внутри документа под каждый случай

        Args:
            cell (_type_): ячейка-объект
            text (str): новый текст
        """
        for paragraph in cell.paragraphs:
            paragraph.text += text
    pass

# ===========================================================================================================

def from_base10_to_baseXX(number_10:int|float, base:int) -> str:
    from decimal import Decimal

    """Перевод числа из десятеричной системы счисления в заданную. Внимание - в дробных числах не гарантируется абсолютная точность!

    Args:
        number_10 (int|float): число в десятеричной системе счисления
        base (int): основание новой системы счисления
    
    Returns:
        str: число в новой системе счисления в формате str. Для дальнейшей работы с этим результатом желательно использовать в связке с классом BaseXX
    """

    def inter_and_fract_parts(number:Decimal):
        """Разделение заданного числа на целую и дробную части

        Args:
            number (Decimal): число в формате Decimal

        Returns:
            tuple[int, Decimal]: целая и дробная части
        """

        in_p = trunc(abs(number))
        fr_p = abs(number) - in_p
        return in_p, fr_p    

    number_10 = Decimal(str(number_10))
    sign, _, e = number_10.as_tuple()
    inter_part, fract_part = inter_and_fract_parts(number_10)

    inter_part_res = []
    if inter_part == 0:
        inter_part_res = ['0']
    else:
        while inter_part > 0:
            inter_part_res.insert(0, ALL_DIGITS[inter_part % base])
            inter_part //= base

    fract_part_res = []
    if e:
        for _ in range(-e + ceil(2 / log10(base))):
            fract_part = fract_part * base
            digit, fract_part = inter_and_fract_parts(fract_part)
            fract_part_res.append(ALL_DIGITS[digit])
    result = f"{'-' if sign else ''}{''.join(inter_part_res)}{';' + ''.join(fract_part_res) if fract_part_res else ''}"
    return result

def from_baseXX_to_base10(number_XX:str, base:int) -> int|float:
    """Перевод числа из заданной системы счисления в десятеричную. Внимание - в дробных числах не гарантируется абсолютная точность!

    Args:
        number_XX (str): число в заданной системе счисления. Передаётся как строка. Предполагается, что эту функцию используют в связке с класом BaseXX
        base (int): основание исходной системы счисления

    Returns:
        int|float: число в десятеричной системе счисления
    """
    if number_XX[0] == '-':
        sign = -1
        number_XX = number_XX[1:]
    else:
        sign = 1
    if ';' not in number_XX:
        number_XX += ';'
    inter_part, fract_part = number_XX.split(';')
    inter_part_res = sum(ALL_DIGITS.index(inter_part[-1 - i]) * pow(base, i) for i in range(len(inter_part)))
    fract_part_res = sum(ALL_DIGITS.index(fract_part[i]) * pow(base, -1 - i) for i in range(len(fract_part))) if fract_part else 0
    result = (inter_part_res + fract_part_res) * sign
    return result

class BaseXX:
    """Класс, представляющий собой число в заданной системе счисления. Поддерживает системы счисления с 2 по 60. Для систем счисления с основанием 2, 8, 10, 16 всё-таки рекомендуется использовать внутренние инструменты. Может работать в двух режимах:

    1. Если в качестве числа задано число (тип int, float), то предполагается, что результат будет число, переведённое из десятеричной системы счисления в заданную
    2. Если в качестве числа задана строка, то предполагается, что это уже число в заданной системе счисления.

    Словом, BaseXX(12, 7) и BaseXX('12', 7) будут двумя разными числами!
    """    

    def __init__(self, number:int|float|str, base:int) -> None:
        if (base < 2) or (60 < base):
            raise ValueError
        self.base = base
        self.digits = ALL_DIGITS[0:self.base]
        if isinstance(number, str):
            self.numberXX = number
            self.number10 = from_baseXX_to_base10(self.numberXX, self.base)
        else:
            self.numberXX = from_base10_to_baseXX(number, self.base)
            self.number10 = number
        pass

    def __str__(self) -> str:
        return f"{self.numberXX} [base{self.base}]"
    
    # def __repr__(self) -> str:
    #     return f"BaseXX({self.number}, {self.base})"
    
    def __add__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 + other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 + Base12(other).number10)
        else:
            return BaseXX(self.number10 + other.number10 if isinstance(other, type(self)) else self.number10 + BaseXX(other, self.base).number10, self.base)
    def __radd__(self, other):
        if isinstance(self, Base12):
            return self + Base12(other)
        else:
            return self + BaseXX(other, self.base)
    def __iadd__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 + other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 + Base12(other).number10)
        else:
            return BaseXX(self.number10 + other.number10 if isinstance(other, type(self)) else self.number10 + BaseXX(other, self.base).number10, self.base)
    
    def __sub__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 - other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 - BaseXX(other).number10)
        else:
            return BaseXX(self.number10 - other.number10 if isinstance(other, type(self)) else self.number10 - BaseXX(other, self.base).number10, self.base)
    def __rsub__(self, other):
        if isinstance(self, Base12):
            return self - Base12(other)
        else:
            return self - BaseXX(other, self.base)
    def __isub__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 - other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 - BaseXX(other).number10)
        else:
            return BaseXX(self.number10 - other.number10 if isinstance(other, type(self)) else self.number10 - BaseXX(other, self.base).number10, self.base)
    
    def __mul__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 * other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 * Base12(other).number10)
        else:
            return BaseXX(self.number10 * other.number10 if isinstance(other, type(self)) else self.number10 * BaseXX(other, self.base).number10, self.base)
    def __rmul__(self, other):
        if isinstance(self, Base12):
            return self * Base12(other)
        else:
            return self * BaseXX(other, self.base)
    def __imul__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 * other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 * Base12(other).number10)
        else:
            return BaseXX(self.number10 * other.number10 if isinstance(other, type(self)) else self.number10 * BaseXX(other, self.base).number10, self.base)
    
    def __truediv__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 / other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 / Base12(other).number10)
        else:
            return BaseXX(self.number10 / other.number10 if isinstance(other, type(self)) else self.number10 / BaseXX(other, self.base).number10, self.base)
    def __rtruediv__(self, other):
        if isinstance(self, Base12):
            return self / Base12(other)
        else:
            return self / BaseXX(other, self.base)
    def __itruediv__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 / other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 / Base12(other).number10)
        else:
            return BaseXX(self.number10 / other.number10 if isinstance(other, type(self)) else self.number10 / BaseXX(other, self.base).number10, self.base)
    
    def __floordiv__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 // other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 // Base12(other).number10)
        else:
            return BaseXX(self.number10 // other.number10 if isinstance(other, type(self)) else self.number10 // BaseXX(other, self.base).number10, self.base)
    def __rfloordiv__(self, other):
        if isinstance(self, Base12):
            return self // Base12(other)
        else:
            return self // BaseXX(other, self.base)
    def __ifloordiv__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 // other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 // Base12(other).number10)
        else:
            return BaseXX(self.number10 // other.number10 if isinstance(other, type(self)) else self.number10 // BaseXX(other, self.base).number10, self.base)        

    def __mod__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 % other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 % Base12(other).number10)
        else:
            return BaseXX(self.number10 % other.number10 if isinstance(other, type(self)) else self.number10 % BaseXX(other, self.base).number10, self.base)
    def __rmod__(self, other):
        if isinstance(self, Base12):
            return self % Base12(other)
        else:
            return self % BaseXX(other, self.base)
    def __imod__(self, other):
        if isinstance(self, Base12):
            return Base12(self.number10 % other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 % Base12(other).number10)
        else:
            return BaseXX(self.number10 % other.number10 if isinstance(other, type(self)) else self.number10 % BaseXX(other, self.base).number10, self.base)
    
    def __lt__(self, other):
        if isinstance(self, Base12):
            return self.number10 < other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 < Base12(other).number10
        else:
            return self.number10 < other.number10 if isinstance(other, type(self)) else self.number10 < BaseXX(other, self.base).number10
    
    def __le__(self, other):
        if isinstance(self, Base12):
            return self.number10 <= other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 <= Base12(other).number10
        else:
            return self.number10 <= other.number10 if isinstance(other, type(self)) else self.number10 <= BaseXX(other, self.base).number10
    
    def __eq__(self, other):
        if isinstance(self, Base12):
            return self.number10 == other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 == Base12(other).number10
        else:
            return self.number10 == other.number10 if isinstance(other, type(self)) else self.number10 == BaseXX(other, self.base).number10
    
    def __ne__(self, other):
        if isinstance(self, Base12):
            return self.number10 != other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 != Base12(other).number10
        else:
            return self.number10 != other.number10 if isinstance(other, type(self)) else self.number10 != BaseXX(other, self.base).number10
    
    def __gt__(self, other):
        if isinstance(self, Base12):
            return self.number10 > other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 > Base12(other).number10
        else:
            return self.number10 > other.number10 if isinstance(other, type(self)) else self.number10 > BaseXX(other, self.base).number10
    
    def __ge__(self, other):
        if isinstance(self, Base12):
            return self.number10 >= other.number10 if isinstance(other, (Base12, BaseXX)) else self.number10 >= Base12(other).number10
        else:
            return self.number10 >= other.number10 if isinstance(other, type(self)) else self.number10 >= BaseXX(other, self.base).number10

class Base12(BaseXX):
    """Класс чисел в 12-ичной системе счисления. Является наследником класса BaseXX, но уже с заданной системой счисления - 12. В качестве цифр для обозначения "10" и "11" используются специальные знаки "↊" и "↋". При задании числа в 12-ичной можно использовать стандартные "A" и "B" (всё равно ни у кого на клавиатуре не будет тех знаков)
    """
    def __init__(self, number:int|float|str) -> None:
        self.base = 12
        if isinstance(number, str) and ('↊' in number or '↋' in number):
            number = number.replace('↊', 'A').replace('↋', 'B')
        super().__init__(number, self.base)
        self.digits = ALL_DIGITS[0:10] + '↊↋'
    
    def __str__(self) -> str:
        return f"{self.numberXX} [base{self.base}]".replace('A', '↊').replace('B', '↋')

def do_something_fun():
    """Код, предложенный Везантиной на запрос "Можешь написать на Питоне код, который делает что-нибудь прикольное?". По её словам, "небольшой код на Python, который создаст анимированный градиентный экран"
    """
    import turtle
    from random import randint

    # настройки экрана
    screen = turtle.Screen()
    screen.setup(700, 700)
    screen.bgcolor("black")

    # настройки черепашки
    turtle.speed(0)
    turtle.penup()

    # создаем цветовую палитру
    colors = ["red", "orange", "yellow", "green", "blue", "purple"]

    # создаем градиентный экран
    for y in range(-350, 350, 10):
        for x in range(-350, 350, 10):
            color = colors[randint(0, len(colors)-1)]
            turtle.goto(x, y)
            turtle.dot(10, color)

    # создаем анимацию
    for _ in range(300):
        turtle.clear()
        for y in range(-350, 350, 10):
            for x in range(-350, 350, 10):
                color = colors[randint(0, len(colors)-1)]
                turtle.goto(x, y)
                turtle.dot(10, color)
            
        # поворот и перемещение черепашки
        turtle.left(5)
        turtle.forward(10)

    turtle.done()

# ===========================================================================================================

if __name__ == '__main__':
    a = BaseXX(5, 7)
    print_debug_mode_on(a % 2 == 0)
    a = a / 2 if a % 2 == 0 else 3*a + 1
    print_debug_mode_on(a)
    a = Base12(18*4*3)
    print(a, a.number10)

    example_dict = {
        'key1': 'value1',
        'key2': {
            'subkey1': 'subvalue1',
            'subkey2': {
                'subsubkey1': 'subsubvalue1'
            }
        },
        'key3': 'value3'
    }
    result_dict = flatten_dictionary(example_dict, True)
    print(result_dict)
    result_dict = flatten_dictionary(example_dict, False)
    print(result_dict)

    # do_something_fun()

    print_debug_mode_on(find_all_systems('V-FC-0001, V-FC-0003, V-FC-0005, V-FC-0011- V-FC-0015'))

    r'\d?[A-Za-zА-Яа-яЁё]\D?\D?\D?\D?\S?\S?\S?\S?\d{1,}[А-Яа-яЁё]?'

    all_test_files = (
        f'_ТЗ_для программы/Веросы/Бланк1/211027853-ОПР АЭРОПРОЕКТ (Реконстр. аэропортового комплекса Чертовицкое г.Воронеж ОАСС).docx',  # 0
        '_ТЗ_для программы/Каналка/В1ап.docx',  # 1
        '_ТЗ_для программы/Вентиляторы/ДУ-1.1.docx',  # 2
        '_ТЗ_для программы/Вентиляторы/неработает/ДУ5-2.docx',  # 3
        '_ТЗ_для программы/Веросы/231037520-ОПР В1.1а СЗ ЖК КИТ (МЖК, г. Мытищи, 25 микрорайон, Шараповский проезд, паркинг).doc',  # 4
        '_ТЗ_для программы/Каналка/ПД5.1 версия 2.docx',  # 5
        '_ТЗ_для программы/Вентиляторы/ПД1,ПД2,ПД3 версия 2.docx',  # 6
        '_ТЗ_для программы/Вентиляторы/ПД5 версия 2.docx',  # 7
        '_ТЗ_для программы/Вентиляторы/ПД4 версия 2.docx',  # 8
        '_ТЗ_для программы/Каналка/П2 (1).docx',  # 9
        '_ТЗ_для программы/Каналка/1SAS01.docx',  # 10
        '_ТЗ_для программы/Каналка/ПВ3А - Таня.docx',  # 11
        '_ТЗ_для программы/Каналка/V-FC-0016 - Регина.docx',  # 12
        '_ТЗ_для программы/Каналка/П7В81 - Лена.docx',  # 13
        '_ТЗ_для программы/Каналка/ПВС-1-20 - 20шт.docx',  # 14
        '_ТЗ_для программы/Веросы/221056572-ОПР ПВ1 КАЛУГАГЛАВСНАБ (Складское помещение, Апрелевка).doc',  # 15
        '_ТЗ_для программы/Веросы/Бланк5-8/1.6. Паровой ТО.docx',  # 16
        '_ТЗ_для программы/Веросы/231019750б-ОПР.doc',  # 17
        '_ТЗ_для программы/Вентиляторы/В6.5.docx',  # 18
        '_ТЗ_для программы/Веросы/231036475-СПБ.doc',  # 19
        '_ТЗ_для программы/Веросы/В таблицинию.doc',  # 20
        '_ТЗ_для программы/Веросы/В лининию.doc',  # 21
        '_ТЗ_для программы/Веросы/231038867-СПБ_ПВ4е.doc',  # 22 - разное вертикальное питерское говно, что-то пошло по пизде
        '_ТЗ_для программы/Веросы/Приточные установки1/231059933-ОПР ТЕПЛОЭЛЕКТРОПРОЕКТ(Новочеркасская ГРЭС)10SAB01.docx',  # 23
        '_ТЗ_для программы/Каналка/В13 (2).docx',  # 24 - тупое говно тупого кАНАЛьного говна говно
        '_ТЗ_для программы/Вентиляторы/МО1.docx',  # 25 - новый бланк вентиляторов
        '_ТЗ_для программы/Веросы/241005620-ОПР 4137.14.16-РС-2; 2А  АСП-Автоматика Газпромнефть-МНПЗ.doc',  # 26
    )

    print(find_all_systems('В1ап'))
    pass

    time_start = perf_counter()
    baba = Blank(all_test_files[26])
    print_debug_mode_on(baba.docx_text)
    print_debug_mode_on.check_debug_mode()
    print_debug_mode_on(baba)
    print_debug_mode_on(repr(baba))
    print_debug_mode_on(baba.blank_type)
    print_debug_mode_on(baba.main_information)
    print_debug_mode_on(baba.all_avaiable_information)
    print_debug_mode_on(*baba.ALL_MAIN_INFO, sep='\n')
    print_debug_mode_on(perf_counter() - time_start)
    print_debug_mode_on(baba.IS_VEROSA)
    pass

    print_debug_mode_on.debug_mode_tumbler()
    print_debug_mode_on.check_debug_mode()

    time_start = perf_counter()
    print_debug_mode_on(baba)
    print_debug_mode_on(repr(baba))
    print_debug_mode_on(baba.blank_type)
    print_debug_mode_on(baba.main_information)
    print_debug_mode_on(baba.all_avaiable_information)
    print_debug_mode_on(*baba.ALL_MAIN_INFO, sep='\n')
    print_debug_mode_on(perf_counter() - time_start)
    pass

    print_debug_mode_on.debug_mode_tumbler()
    print_debug_mode_on.check_debug_mode()

    time_start = perf_counter()
    print_debug_mode_on(baba)
    print_debug_mode_on(repr(baba))
    print_debug_mode_on(baba.blank_type)
    print_debug_mode_on(baba.main_information)
    print_debug_mode_on(baba.all_avaiable_information)
    print_debug_mode_on(*baba.ALL_MAIN_INFO, sep='\n')
    print_debug_mode_on(perf_counter() - time_start)
    pass

    for test_file in all_test_files:
    # test_file = all_test_files[43]
        print(test_file)
        dodoco = Document(test_file)
        test_blank = Blank(dodoco)
        print(test_blank, test_blank.all_main_info_text, sep='\n')
        print(test_blank.blank_type)

        print(find_all_systems(test_blank.main_information['Название']))
else:
    print_debug_mode_on.debug_mode_tumbler()